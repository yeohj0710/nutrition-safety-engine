#!/usr/bin/env python3
"""방법론 중심 v2 논문의 구조·수치·경계를 검증한다."""
import hashlib,json,re
from pathlib import Path
from docx import Document
from pypdf import PdfReader
R=Path(__file__).resolve().parents[1]; D=R/"research/thesis/ai_exploratory_thesis.docx"; P=R/"research/thesis/ai_exploratory_thesis.pdf"; M=R/"research/thesis/ai_exploratory_thesis_ko.md"
def sha(p):return hashlib.sha256(p.read_bytes()).hexdigest()
def main():
 e=[]; doc=Document(D); pdf=PdfReader(P); text="\n".join(x.text for x in doc.paragraphs)+"\n"+"\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells); ptext="\n".join(x.extract_text() or "" for x in pdf.pages)
 required=["1. 서론","2. 이론적·방법론적 배경","3. 연구 방법","3.17 설계 대안과 선택 근거","4. 연구 결과","5. 고찰","6. 연구의 한계","7. 결론","20,230","18,015","2,215","12,330","6,649","임상행동 누출","사람의 독립 선별"]
 def compact(value): return re.sub(r"\s+", "", value)
 doc_compact,pdf_compact=compact(text),compact(ptext)
 for x in required:
  if compact(x) not in doc_compact or compact(x) not in pdf_compact:e.append("missing:"+x)
 if len(doc.paragraphs)<170:e.append("paragraphs<170")
 if len(doc.tables)<5:e.append("tables<5")
 if len(pdf.pages)<12:e.append("pdf_pages<12")
 if any(x in text for x in ["TODO","TBD","사람 선별 완료","체계적 문헌고찰 결과"]):e.append("forbidden expression")
 out={"status":"valid" if not e else "invalid","errors":e,"docx_paragraphs":len(doc.paragraphs),"docx_tables":len(doc.tables),"pdf_pages":len(pdf.pages),"pdf_text_chars":len(ptext),"docx_sha256":sha(D),"pdf_sha256":sha(P),"markdown_sha256":sha(M)}; print(json.dumps(out,ensure_ascii=False,indent=2)); return bool(e)
if __name__=="__main__":raise SystemExit(main())
