#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from pypdf import PdfReader
import csv,json,hashlib,re,zipfile
R=Path(__file__).resolve().parents[1];d=R/"research/thesis/systematic_review_thesis_v3.docx";p=R/"research/thesis/systematic_review_thesis_v3.pdf";m=json.loads((R/"research/systematic_review_v3/manifest.json").read_text(encoding="utf-8"));c=json.loads((R/"research/systematic_review_v3/core_manifest.json").read_text(encoding="utf-8"));q=json.loads((R/"research/validation/software_quality_v3.json").read_text(encoding="utf-8"))
doc=Document(d);pdf=PdfReader(p);text="\n".join(x.text for x in doc.paragraphs)+"\n"+"\n".join(x.extract_text() or "" for x in pdf.pages);errors=[]
with (R/"research/systematic_review_v3/core_evidence.csv").open(encoding="utf-8",newline="") as handle:first_reference=next(csv.DictReader(handle))
for x in ["체계적 근거 검토","개인맞춤 조회 시스템",f"{m['records']:,}",f"{m['with_dose']:,}",f"{m['with_fulltext_locator']:,}",str(c['core_records']),"https://nutrition-safety-engine.vercel.app",f"{q['tests']}개 자동 테스트"]:
 if x not in text:errors.append("missing:"+x)
expected_a2=f"오메가-3와 항응고제 질문은 {c['per_question']['A2']}건"
if expected_a2 not in text:errors.append("missing dynamic A2 count:"+expected_a2)
for x in ["Vercel Preview 환경","53개 자동 테스트","AI 기반 탐색적 근거지도 구축"]:
 if x in text:errors.append("stale:"+x)
if len(pdf.pages)<15:errors.append("pdf_pages<15")
if len(doc.tables)<4:errors.append("tables<4")
if re.search(r"\b(?:19|20)\d{2}\.0\b",text):errors.append("float-formatted bibliography year")
if re.search(r"보고 용량:[^\n]{0,160}(?:(?:^|\| )000 IU|\b\d+-\d+\s+G(?:\b|$)|mg/dL)",text):errors.append("malformed or non-dose unit in thesis")
if first_reference["authors"].split(";")[0].strip() not in text or first_reference["venue"] not in text:errors.append("bibliography author or venue missing")
with zipfile.ZipFile(d) as archive:
 styles=archive.read("word/styles.xml").decode("utf-8")
 document_xml=archive.read("word/document.xml").decode("utf-8")
if "Pretendard" not in styles or "Pretendard" not in document_xml:errors.append("Pretendard missing from DOCX XML")
if "Malgun Gothic" in styles or "맑은 고딕" in styles:errors.append("legacy Malgun Gothic remains in styles.xml")
print(json.dumps({"status":"valid" if not errors else "invalid","errors":errors,"pages":len(pdf.pages),"paragraphs":len(doc.paragraphs),"tables":len(doc.tables),"docx_sha256":hashlib.sha256(d.read_bytes()).hexdigest(),"pdf_sha256":hashlib.sha256(p.read_bytes()).hexdigest()},ensure_ascii=False,indent=2));raise SystemExit(bool(errors))
