#!/usr/bin/env python3
"""Build a visibly non-final Korean methods checkpoint DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
SOURCE = REPO / "research/thesis/checkpoints/methods_checkpoint_nonfinal.md"
OUTPUT = REPO / "research/thesis/checkpoints/methods_checkpoint_nonfinal.docx"


def set_font(run, name: str, size: float, bold: bool | None = None, color: str = "202124"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_font(style, name: str, size: float, bold: bool = False, color: str = "202124"):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (fld_char, instr, separate, text, end):
        run._r.append(node)
    set_font(run, "맑은 고딕", 8.5, color="6B7280")


def add_bottom_border(paragraph, color: str = "D1D5DB", size: str = "6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    set_repeat_font(normal, "바탕", 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.55
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.first_line_indent = Cm(0.8)
    normal.paragraph_format.keep_together = False

    for name, size, before, after, color in (
        ("Heading 1", 17, 18, 10, "111827"),
        ("Heading 2", 13, 14, 7, "1F3A5F"),
        ("Heading 3", 11.5, 10, 5, "374151"),
    ):
        style = styles[name]
        set_repeat_font(style, "맑은 고딕", size, True, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = None

    if "Checkpoint Note" not in styles:
        note = styles.add_style("Checkpoint Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Checkpoint Note"]
    set_repeat_font(note, "맑은 고딕", 9.5, False, "4B5563")
    note.paragraph_format.left_indent = Cm(0.5)
    note.paragraph_format.right_indent = Cm(0.5)
    note.paragraph_format.space_before = Pt(5)
    note.paragraph_format.space_after = Pt(10)
    note.paragraph_format.line_spacing = 1.35
    note.paragraph_format.first_line_indent = None


def configure_section(section, first: bool = False):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(23)
    section.left_margin = Mm(28)
    section.right_margin = Mm(28)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)
    section.different_first_page_header_footer = first


def add_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "방법론 체크포인트 · 비최종본"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(3)
    set_font(p.runs[0], "맑은 고딕", 8, color="6B7280")
    add_bottom_border(p, "D1D5DB", "4")
    add_page_number(section.footer.paragraphs[0])


def blocks(markdown: str):
    lines = markdown.splitlines()
    current: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#") or stripped == "\\pagebreak":
            if current:
                yield "paragraph", " ".join(part.strip() for part in current).strip()
                current = []
            if stripped == "\\pagebreak":
                yield "pagebreak", ""
            else:
                level = len(stripped) - len(stripped.lstrip("#"))
                yield f"h{level}", stripped[level:].strip()
        elif not stripped:
            if current:
                yield "paragraph", " ".join(part.strip() for part in current).strip()
                current = []
        else:
            current.append(line.replace("  ", "\n"))
    if current:
        yield "paragraph", " ".join(part.strip() for part in current).strip()


def main() -> int:
    content = list(blocks(SOURCE.read_text(encoding="utf-8")))
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], first=True)
    doc.core_properties.title = "영양보충제 안전성 연구 방법론 체크포인트"
    doc.core_properties.subject = "결과 동결 전 비최종 연구방법 문서"
    doc.core_properties.author = "여형준"

    # Cover, intentionally plain and clearly non-final.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(70)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run(content[0][1])
    set_font(run, "맑은 고딕", 19, True, "111827")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(subtitle.add_run(content[1][1]), "맑은 고딕", 13, True, "1F3A5F")

    notice = doc.add_paragraph(style="Checkpoint Note")
    notice.alignment = WD_ALIGN_PARAGRAPH.LEFT
    notice.paragraph_format.space_before = Pt(14)
    notice.paragraph_format.space_after = Pt(18)
    notice.add_run(content[2][1])
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F4F6")
    notice._p.get_or_add_pPr().append(shading)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.line_spacing = 1.5
    meta.paragraph_format.space_before = Pt(10)
    for line in content[3][1].split("\n"):
        set_font(meta.add_run(line.strip()), "맑은 고딕", 10, color="4B5563")
        meta.add_run().add_break()

    # Skip cover blocks through explicit page break.
    start = next(i for i, block in enumerate(content) if block[0] == "pagebreak") + 1
    doc.add_page_break()
    add_running_furniture(doc.sections[0])
    for kind, value in content[start:]:
        if kind == "pagebreak":
            doc.add_page_break()
        elif kind == "h1":
            doc.add_paragraph(value, style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(value, style="Heading 2")
        elif kind == "h3":
            doc.add_paragraph(value, style="Heading 3")
        elif kind == "paragraph":
            p = doc.add_paragraph(value)
            p.paragraph_format.widow_control = True

    # Mark the final hold section without implying a result.
    for paragraph in doc.paragraphs:
        if paragraph.text == "작성 보류 항목":
            paragraph.paragraph_format.page_break_before = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
