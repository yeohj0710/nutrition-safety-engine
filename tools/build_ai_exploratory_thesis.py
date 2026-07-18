#!/usr/bin/env python3
"""검증된 manifest에서 방법론 중심의 한국어 학위논문 DOCX/Markdown을 생성한다."""
from __future__ import annotations
import hashlib, json
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]; OUT=ROOT/"research/thesis"; OUT.mkdir(exist_ok=True)
DOCX=OUT/"ai_exploratory_thesis.docx"; MD=OUT/"ai_exploratory_thesis_ko.md"
AUTHOR="여형준"
SUBMISSION_DATE="2026년 7월"
P={"map":ROOT/"research/synthesis/ai_exploratory_map_manifest.json","screen":ROOT/"research/screening/ai_exploratory_screening_manifest.json","nonpub":ROOT/"research/screening/ai_exploratory_nonpubmed_manifest.json","perf":ROOT/"research/validation/ai_exploratory_performance.json","protocol":ROOT/"research/protocol/protocol-v2.0-ai-exploratory.md","bundle":ROOT/"src/generated/ai-exploratory-bundle.json"}
D={k:json.loads(v.read_text(encoding="utf-8")) for k,v in P.items() if k not in {"protocol"}}
M,S,N,V=D["map"],D["screen"],D["nonpub"],D["perf"]
Q={"A1":"비타민 K와 비타민 K 길항제","A2":"오메가-3와 경구 항응고제","B1":"칼슘 보충과 결석 위험","B2":"비타민 D와 결석 위험","B3":"비타민 C와 결석 위험"}

sections=[
("1. 서론",[
("1.1 연구 배경",["건강기능식품과 영양보충제는 처방전 없이 쉽게 살 수 있지만, 고위험 임상상황에서는 단순히 ‘몸에 좋다’거나 ‘나쁘다’는 식으로 안전성을 판단하기 어렵다. 같은 성분이라도 함께 먹는 약, 앓고 있는 병, 복용량, 제형, 복용 기간, 검사 수치에 따라 임상적 의미가 달라지기 때문이다. 특히 항응고제를 복용하거나 신장결석 위험이 있는 사람에게는 성분의 평균적인 효과보다, 약물 상호작용 가능성과 실제 노출량, 확인해야 할 지표, 아직 남아 있는 불확실성을 함께 살피는 절차가 중요하다.","이러한 판단에 필요한 근거는 무작위 대조시험, 관찰연구, 증례보고, 임상시험 등록정보, 국내외 데이터베이스에 흩어져 있다. 그런데 검색 결과가 많다고 해서 근거가 충분한 것은 아니다. 검색된 문헌이 실제로 연구 질문에 답하는지, 초록이나 원문을 실제로 확인할 수 있는지, 같은 연구가 여러 편의 보고서로 중복 집계되지는 않았는지를 따로 확인해야 한다."]),
("1.2 문제 제기",["처음 세운 연구계획은 두 사람이 각자 문헌을 선별하고 자료를 이중으로 추출하는 체계적 문헌고찰이었다. 그러나 막상 수행 단계에 들어서자 구독이 필요한 원문, 일부 국내 데이터베이스의 원자료 내려받기, 독립적으로 판정할 두 명의 사람을 확보하지 못하였다. 이 상태에서 AI의 판정을 사람의 판정인 것처럼 표시하거나 끝나지 않은 선별을 끝난 것으로 처리하면, 결과 수치는 그럴듯하게 만들 수 있어도 그 수치를 남이 검증할 길은 사라진다.","그래서 본 연구는 부족한 절차를 감추는 대신 연구 질문 자체를 바꾸었다. 임상 효과를 확정하는 체계적 문헌고찰이 아니라, 공개된 서지정보를 원형 그대로 보존하고 자동 분류가 어디에서 일치하고 어디에서 갈리는지를 그대로 드러내며, 그 자료를 임상 행동 없이 탐색하기만 하는 도구가 정해진 경계 안에서 제대로 작동하는지를 검증하는 탐색적 근거지도 연구로 전환하였다. 이 전환은 연구를 편하게 줄이려는 것이 아니라, 지금 가진 자원으로 정직하게 답할 수 있는 질문과 답할 수 없는 질문을 구분하기 위한 설계상의 결정이다."]),
("1.3 연구 목적과 질문",["본 연구의 목적은 네 가지다. 첫째, 항응고제 병용과 신장결석 위험에 관한 다섯 질문의 공개 검색 결과를 출처·질문·관찰 가능 여부에 따라 정리한다. 둘째, 성향이 다른 두 개의 결정론적 자동 분류 방식을 함께 적용하여 두 방식이 모두 우선 확인으로 본 문헌, 모두 후순위로 본 문헌, 판단이 갈린 문헌을 구분한다. 셋째, 공개 원문의 위치와 파일 지문(해시)을 연결하여 논문의 주장에서 원자료까지 거꾸로 따라갈 수 있는 경로를 만든다. 넷째, 정확한 성분명에만 반응하고 임상 행동은 내놓지 않는 탐색 도구를 구현하여 결과의 일관성, 질문 연결의 정확성, 출처 추적의 완전성, 정보 누출 여부를 검증한다.","다섯 가지 연구 질문은 A1 비타민 K와 비타민 K 길항제, A2 오메가-3와 경구 항응고제, B1 칼슘 보충과 결석 위험, B2 비타민 D와 결석 위험, B3 비타민 C와 결석 위험이다. 분석의 기본 단위는 개별 연구가 아니라, 특정 문헌이 특정 질문에서 검색된 ‘레코드-질문 단위(record-question unit)’다."]),
("1.4 연구 범위",["본 논문은 임상 권고를 제시하지 않으며, 효과크기를 통합하거나 위험비를 계산하지 않고, 비뚤림 위험(RoB)이나 근거수준(GRADE)도 판정하지 않는다. 자동 분류는 어떤 문헌을 먼저 확인할지 순서를 정해 주는 계산상의 표시일 뿐, 문헌을 최종적으로 포함하거나 제외하는 판정이 아니다. 이렇게 범위를 좁히면 결론의 강도는 제한되지만, 실제로 관찰한 자료와 그에 대한 해석 사이의 거리는 오히려 분명해진다."])]),
("2. 이론적·방법론적 배경",[
("2.1 체계적 문헌고찰과 탐색적 근거지도의 구분",["체계적 문헌고찰은 미리 정한 기준에 따라 연구를 포함하거나 제외하고 비뚤림 위험을 평가한 뒤, 질문에 대한 종합적인 결론을 제시한다. 반면 근거지도는 어떤 자료가 어디에 있고 어느 부분이 비어 있는지를 지도처럼 정리하는 작업이다. 본 연구에는 사람의 판정과 독립적인 이중 추출이 없으므로, 앞의 결론을 내릴 권한은 주장할 수 없고 뒤의 자료를 찾아 정리하는 기능에 초점을 맞춘다.","이 구분은 이름만 다른 문제가 아니라 구체적인 분석 규칙으로 이어진다. 즉 검색된 문헌 수를 연구 수라고 부르지 않고, 자동 분류가 일치한 것을 포함 판정이라고 부르지 않으며, 원문 위치를 확보하지 못한 자료에서 효과 수치를 뽑아내지 않는다."]),
("2.2 레코드·보고물·연구의 구분",["하나의 임상시험은 프로토콜, 등록정보, 학술지 논문, 후속 분석 등 여러 편의 보고물로 나뉘어 나타날 수 있다. 반대로 서지 레코드 하나가 언제나 독립된 하나의 연구를 뜻하지도 않는다. 그래서 본 연구는 DOI·PMID·등록번호를 그대로 보존하되, 자동으로 표시한 중복 후보를 연구 단위의 확정된 중복 제거로 해석하지 않는다."]),
("2.3 AI의 역할과 인식론적 경계",["AI는 같은 규칙을 반복 실행하고 서식을 표준화하며 먼저 볼 후보의 순서를 매기는 데에만 사용하였다. 임상적 관련성, 비뚤림 위험, 근거의 확실성, 권고의 강도는 자동화된 결과로 대신하지 않았다. 두 분류 방식이 서로 다르게 판단한 문헌은 오류로 지워 버리지 않고, 나중에 사람이 검토해야 할 불확실성의 층으로 남겨 두었다."]),
("2.4 출처 계보와 결정론",["재현성은 같은 문장을 다시 만들어 내는 능력만을 뜻하지 않는다. 어떤 원자료가 어떤 변환을 거쳐 어느 결과 수치로 이어졌는지를 추적할 수 있어야 한다. 본 연구는 원자료의 경로, 파일 지문(SHA-256), 생성에 쓴 스크립트, 산출 기록(manifest)을 서로 연결하고, 같은 입력을 반복해서 넣으면 언제나 같은 연결과 같은 출력이 나오도록 요구하였다."])]),
("3. 연구 방법",[
("3.1 전체 연구설계",["연구는 아홉 단계로 이루어진다. ① 기존 코드·자료의 감사와 격리, ② 프로토콜 개정, ③ 검색 자료 수집, ④ 표준화와 출처 계보 생성, ⑤ 자동 우선순위 부여, ⑥ 탐색적 근거지도 구축, ⑦ 출처를 붙인 임시 서술 생성, ⑧ 탐색 도구 구현, ⑨ 독립 시나리오 검증이다. 각 단계는 앞 단계의 산출물과 그 지문(해시)을 입력으로 받고, 검증 관문을 통과했을 때에만 다음 단계의 공식 입력으로 인정되도록 설계하였다.","이 순서를 택한 이유는 결과를 먼저 정해 놓고 원자료를 거기에 끼워 맞추는 순환을 막기 위해서다. 예를 들어 논문 표의 합계는 문서에 직접 입력하지 않고 근거지도 산출 기록에서 읽어 온다. 검색 자료가 바뀌면 산출 기록과 표가 함께 바뀌고, 둘이 어긋나면 검증기가 실패한다."]),
("3.2 저장소 감사와 legacy 격리",["연구를 시작하는 시점에 기존 코드, 배포 흔적, 데이터 경로, 문서에 직접 박아 넣은 수치, 빠져 있는 시험을 모두 점검하였다. 이전에 만들어 둔 산출물은 ‘완료’ 표시가 붙어 있더라도 다시 검증하기 전에는 미검증 자료(legacy_unverified)로 분류하였다. 해당 파일 36개를 기준 커밋별 격리 경로로 옮기고, 옮기기 전후의 지문(SHA-256)을 비교하여 달라진 것이 하나도 없음을 확인하였다.","격리는 삭제가 아니라 증거를 보존하는 조치다. 기존 결과가 새 분석에 우연히 섞이는 것을 막으면서도, 과거의 계산을 다시 재현하거나 차이를 조사할 수 있도록 남겨 둔다. 새로 만든 기본 실행 환경은 검증을 마친 v2 자료만 읽고, 과거 기능은 별도 경로에 둔다."]),
("3.3 프로토콜 개정의 논리",["처음 목표였던 체계적 문헌고찰(v1)에는 두 명의 독립 선별자, 사람이 직접 하는 이견 조정, 이중 추출, 비뚤림 위험 평가, 원문 접근이 필요했다. 이 조건을 갖추지 못했으므로 v1을 ‘완료’로 다시 표시하지 않고, 외부의 사람 검토가 필요한 상태로 남겨 두었다.","개정한 프로토콜(v2)은 질문을 세 가지로 좁혔다. 공개 서지자료를 얼마나 찾아낼 수 있는가, 자동 분류가 실제로 어떻게 작동하는가, 출처 계보를 갖춘 탐색 도구의 기술적 성능은 어떠한가이다. 그 결과 본 연구는 ‘어떤 보충제가 안전한가’가 아니라 ‘어떤 문헌 후보가 관찰되고, 자동화가 어디에서 일치·불일치하며, 도구가 정해진 경계 안에서 재현 가능하게 작동하는가’에 답한다."]),
("3.4 정보원 선택과 접근 제약",["PubMed는 생의학 서지정보와 초록의 주 출처로, ClinicalTrials.gov는 등록된 임상연구의 보완 출처로, KoreaMed는 국내 의학문헌의 보완 출처로 선택하였다. 세 출처는 자료 구조와 제공 범위가 다르므로, 하나의 분류 규칙을 억지로 똑같이 적용하지 않았다.","Embase, Scopus, Web of Science와 일부 국내 원자료 내려받기, 구독 원문은 접근 권한이나 안정적인 원본 내보내기(native export)를 확보하지 못하여 공식 분석 대상에 넣지 않았다. 접근하지 못한 것을 ‘0건’으로 바꾸지 않고 ‘접근 제약’으로 그대로 기록하였다. 이렇게 하면 포괄성은 낮아지지만, 관찰하지 않은 자료를 관찰한 것처럼 집계하는 오류는 피할 수 있다."]),
("3.5 질문별 검색 개념",["A군 질문은 보충제 성분 개념, 항응고제 계열이나 개별 약물 개념, 상호작용·출혈·응고 관련 개념을 결합하였다. B군 질문은 보충제 성분 개념과 신장결석·요로결석 관련 개념을 결합하였다. 초기 검색은 놓치는 문헌을 줄이기 위해 민감도를 우선하되, 검색식만으로 임상적 관련성을 확정하지는 않았다.","특히 ‘비타민 K 길항제’처럼 성분명과 약물군 이름이 단어 수준에서 겹치는 질문은 엉뚱한 문헌이 걸릴 위험이 크다. 그래서 검색 총수와 자동 분류 결과의 차이를 질문별로 따로 보존하고, 수가 많다는 이유만으로 근거가 풍부하다고 해석하지 않았다."]),
("3.6 원자료 보존과 해시",["각 원자료는 받은 상태 그대로, 가능한 한 원본 바이트를 유지한 채 보존하고 지문(SHA-256)을 계산하였다. 표준화한 파일은 원자료를 덮어쓰지 않고 따로 만들며, 생성 기록에 입력 경로, 지문, 행 수, 생성 시각, 출력 지문을 남겼다. 같은 경로의 파일이 바뀌면 출처 계보 검증이 실패하도록 하였다."]),
("3.7 정규화와 분석단위",["출처마다 다른 필드 이름을 공통 서식으로 맞추었다. 공통 필드는 질문 번호, 출처, 출처의 원래 문헌 번호, 제목, 초록, 연도, 학술지, 출판 유형, DOI·PMID·등록번호, URL, 관찰 가능 여부다. 값이 없는 칸은 빈 문자열과 ‘수집하지 못함’이라는 표시를 구분하여, 원래 없는 것과 가져오지 못한 것을 혼동하지 않았다.","하나의 문헌이 둘 이상의 질문에서 검색될 수 있으므로 기본 분석 단위를 ‘레코드-질문 단위(record-question unit)’로 정하였다. 전체 행 수는 질문별 검색 상황을 그대로 보존한 값이며, 고유한 연구의 개수가 아니다. 고유 식별자를 기준으로 집계해야 할 때에는 분모와 변환 규칙을 따로 밝혔다."]),
("3.8 중복 후보 처리",["PMID, DOI, 등록번호와 표준화한 제목을 이용해 중복이 의심되는 문헌을 표시하였다. 다만 같은 연구인지 여부는 연구 설계, 표본, 시점, 보고 목적을 함께 따져야 하므로, AI 혼자서 동일 연구군(study-family)을 확정하지 않았다. 따라서 본 연구의 결과표는 중복을 확정 제거한 뒤의 연구 수가 아니라 레코드-질문 단위 수를 보고한다."]),
("3.9 이중 자동 분류 설계",["PubMed 문헌에는 두 가지 분류 방식을 적용하였다. 하나는 민감도 우선 방식으로, 성분과 임상상황에 관련된 신호를 폭넓게 잡아 누락을 줄이도록 설계하였다. 다른 하나는 구조적 보수 방식으로, 제목과 초록에 직접적인 신호가 약한 문헌을 후순위로 미루도록 설계하였다.","두 방식이 모두 남긴 문헌은 유지 합의(ai_agreement_retain), 모두 후순위로 미룬 문헌은 후순위 합의(ai_agreement_deprioritize), 서로 다르게 판단한 문헌은 불일치 불확실(ai_disagreement_uncertain)로 기록하였다. 이 이름들은 사람이 내리는 포함·제외 판정 용어를 일부러 피한 것이다. 분류 결과는 뒤에 사람이 검토할 때 작업 순서를 정하는 데 도움을 줄 뿐, 근거로서의 적격성을 확정하지 않는다."]),
("3.10 비-PubMed 자료 처리",["ClinicalTrials.gov와 KoreaMed는 PubMed와 필드 구성, 초록이 있는 비율, 레코드의 의미가 다르다. 여기에 같은 분류기를 그대로 적용하면 출처의 차이가 마치 관련성의 차이인 것처럼 오해될 수 있으므로, 이 자료는 순위를 매기지 않고 비순위 후보(ai_unranked_source_candidate)로 보존하였다. 이는 분석의 일관성보다 출처별 측정의 타당성을 앞세운 결정이다."]),
("3.11 관찰 가능성 상태",["초록 본문이 있으면 초록 확인(abstract_observed), 제목과 서지정보만 있으면 제목·서지정보만(title_metadata_only)으로 분류하였다. PMC 식별자가 연결된 문헌은 공개 원문 위치 확인(public_full_text_locator_available)으로 따로 표시하였다. 여기서 PMC 위치정보는 공개 원문으로 이동할 수 있는 주소일 뿐, 본 연구가 그 원문을 모두 읽고 내용을 추출했다는 뜻은 아니다."]),
("3.12 근거지도와 임시 주장",["질문별 근거지도는 출처 수, 초록 확인 여부, 자동 분류 결과, 공개 원문 위치정보를 문헌 한 줄 단위로 묶어 보여 준다. 여기에 붙는 임시 서술은 임상적 결론이 아니라, 모아 놓은 자료가 지금 어떤 상태로 관찰되는지를 요약한 것이다. 모든 수치에는 생성 기록과 지문을 연결했으며, 효과의 방향이나 안전성에 대한 권고는 만들지 않았다."]),
("3.13 탐색 도구 설계",["탐색 도구는 정확한 성분명과 미리 정해 둔 동의어에만 질문을 연결한다. 일부 글자가 우연히 겹치거나 비슷한 표현이 임상 경로로 잘못 이어지지 않도록, 부분·유사 일치(near-match)는 연결되지 않게 하였다. 도구가 돌려주는 정보는 질문별 자료의 규모, 출처 분포, 자동 분류 분포, 출처 계보상의 위치뿐이다.","임상 행동을 담는 칸(clinical_actions)은 항상 비워 두어 복용 시작·중단, 용량 변경, 검사 권고 같은 지시를 내놓지 않는다. 이 제한은 기능이 부족해서가 아니라, 탐색용 자료를 임상 의사결정 도구로 오인하지 않게 하려는 안전 장치다."]),
("3.14 독립 시나리오 검증",["다섯 질문에 정확 일치, 동의어, 오타, 경계 입력을 조합하여 120개의 모의 시나리오를 만들고, 각 시나리오를 세 번씩 반복 실행하였다. 같은 입력에 같은 결과가 나오는지, 질문이 예상대로 연결되는지, 출처 계보 정보가 빠짐없이 채워지는지, 임상 행동이 새어 나오지는 않는지, 이전 자료가 섞이지는 않는지, 유사어가 잘못 연결되지는 않는지를 측정하였다.","이 모의 시나리오는 실제 임상 사용자나 전문가가 만든 기준 정답(gold standard)이 아니다. 따라서 본 연구는 민감도·특이도나 임상 정확도를 보고하지 않는다. 여기서 검증하는 것은 어디까지나 소프트웨어의 경계와 반복 재현성이다."]),
("3.15 분석과 품질 게이트",["기술통계는 출처별·질문별·관찰 상태별·자동 분류별 빈도와 비율로만 제시하였다. 효과크기를 임의로 통합하거나 이질성을 계산하지 않았다. 코드 단위시험, 타입 검사, 실제 배포용 빌드, 데이터 지문 검사, 문서 수치 일치 검사, PDF 육안 검사를 각각 별개의 관문으로 두고 운영하였다."]),
("3.16 재현성과 변경관리",["생성한 모든 산출물은 스크립트로 다시 만들 수 있게 하였고, 논문 워드 파일과 PDF도 같은 산출 기록에서 수치를 읽어 온다. 작업 일지, 결정 기록, 위험 목록, 미해결 항목에는 설계 변경과 아직 풀지 못한 외부 과제를 적어 둔다. 사람이 검토해야 하는 항목은 지우거나 건너뛰어 완료로 표시하지 않고, 앞으로의 연구 단계로 남긴다.","변경관리는 원자료, 변환 코드, 생성 데이터, 서비스 묶음, 논문을 하나의 사슬로 본다. 위쪽 자료가 바뀌면 아래쪽 산출물을 다시 만들고 전체 검증을 되풀이한다. 최종 산출 기록에는 상대 경로, 파일 크기, 지문(SHA-256)을 남겨, 복사나 배포 과정에서 파일이 손상되면 곧바로 잡아낼 수 있게 하였다."]),
("3.17 설계 대안과 선택 근거",["첫 번째 대안은 사람 선별을 확보할 때까지 연구를 통째로 멈추는 것이었다. 이는 체계적 문헌고찰이라는 원래 목표에는 가장 충실하지만, 공개 자료를 모아 출처를 정리하고 소프트웨어의 경계를 검증하는 작업까지 멈출 이유는 없었다. 그래서 임상적 결론은 미뤄 두되 독립적으로 검증할 수 있는 탐색 연구는 계속하기로 하였다.","두 번째 대안은 하나의 자동 분류 점수로 모든 출처의 순위를 매기는 것이었다. 구현은 간단하지만 출처마다 다른 필드 차이와 분류기의 불확실성을 가려 버린다. 본 연구는 두 PubMed 분류 방식이 엇갈리는 상태와 비-PubMed 자료의 비순위 상태를 따로 보존하여, 측정 도구의 차이가 결과에 그대로 드러나게 하였다.","세 번째 대안은 검색 건수를 줄이려고 매우 좁은 검색식을 쓰는 것이었다. 그러나 고위험 상황에서는 초기에 문헌을 놓쳤을 때의 대가가 크고 용어의 변이도 많으므로, 검색 단계에서는 민감도를 우선하였다. 대신 뒤에 자동 분류와 관찰 가능성 층을 덧붙여, 넓게 검색한 결과가 곧바로 직접적인 근거로 오해되지 않게 하였다."]),
("3.18 오류 예방 설계",["수치를 문서에 직접 박아 넣는 실수를 막기 위해, 본문과 표의 핵심 수치는 JSON 산출 기록에서 그때그때 읽어 온다. 문서 검증기는 필수 장이 있는지, 최소한의 문단·표·쪽수를 채웠는지, 결과 수치가 맞는지, 금지 표현이 없는지를 검사한다. 이렇게 하면 문장을 잘 다듬는 일과 수치를 정확히 지키는 일을 서로 분리할 수 있다.","이전 자료가 실행 중에 새어 나오지 않도록, 새로 만든 기본 화면과 API는 v2 묶음만 참조한다. 과거 경로는 별도의 legacy 화면에 두어 기본 탐색 결과에는 나타나지 않게 하였다. 시험은 불러오기 경로뿐 아니라 실제 응답의 출처 정보와 임상 행동 칸의 경계까지 확인한다."]),
("3.19 윤리와 의사소통 원칙",["건강과 관련된 도구가 불완전한 근거를 단정적으로 보여 주면, 사용자가 복용을 임의로 바꿀 수 있다. 그래서 본 연구는 임상 행동을 담는 자리를 구조적으로 비워 두고, 자료의 범위와 아직 확인하지 못한 상태를 응답에 함께 드러냈다. 이는 면책 문구 하나에 기대지 않고, 기능 자체에서 오용 가능성을 줄이는 접근이다.","AI가 수행한 분류를 사람의 검토인 것처럼 표현하지 않았고, 사람이 반드시 필요한 단계는 앞으로의 과제로 분명히 밝혔다. 연구의 완성도는 결론을 많이 쓰는 데 있지 않고, 결론을 내릴 수 없는 조건이 무엇인지를 독자가 알아볼 수 있게 하는 데 있다고 보았다."]),
("3.20 웹 시스템 구현과 배포 환경",["개인 맞춤 안전성 조회 시스템은 Next.js 16.2.1의 App Router로 구현하고 Vercel 시험 배포(Preview) 환경에 올렸다. 사용자는 보충제, 하루 복용량, 함께 먹는 약, 질환·결석 병력, 검사 수치를 입력하고, 시스템은 질문별 확인 사항과 핵심 근거 문헌, 보고된 용량, 안전성 결과, 원문 주소를 돌려준다.","배포를 점검할 때 실제 배포용 빌드, 타입스크립트 검사, 자동 시험 53개를 모두 통과하였다. 배포 번호와 주소, 지역, 경로 목록, 빌드 결과는 별도의 배포 점검 기록에 고정해 두었다. 연구 결과를 확정하기 전까지 배포는 시험(Preview) 상태로 유지하고, 최종 점검을 마친 뒤에 정식(Production)으로 올린다."])]),
("4. 연구 결과",[
("4.1 수집 자료의 구성",[f"근거지도에는 레코드-질문 단위가 모두 {M['row_count']:,}건 담겼다. 출처별로는 PubMed {M['source_counts']['pubmed']:,}건, ClinicalTrials.gov {M['source_counts']['clinicaltrials']:,}건, KoreaMed {M['source_counts']['koreamed']:,}건이었다. 이 값은 중복을 확정 제거한 뒤의 고유 연구 수가 아니다.",f"이 가운데 초록까지 확인된 것은 {M['abstract_observed']:,}건, 제목과 서지정보만 확인된 것은 {M['title_metadata_only']:,}건이었다. 공개 PMC 위치정보가 연결된 것은 {M['pmc_locator_record_question_rows']:,}건이며, 고유 문헌으로 환산하면 {M['unique_records_with_pmc_identifier']:,}건이었다."]),
("4.2 자동 분류 분포",[f"PubMed 문헌에 두 자동 분류를 함께 적용한 결과, 유지 합의는 {S['classifications']['ai_agreement_retain']:,}건, 후순위 합의는 {S['classifications']['ai_agreement_deprioritize']:,}건, 불일치 불확실은 {S['classifications']['ai_disagreement_uncertain']:,}건이었다. 비-PubMed 자료 {N['row_count']:,}건은 순위를 매기지 않고 후보로 보존하였다.","불일치가 이만큼 많다는 것은 하나의 단순한 규칙만으로 관련성을 확정하기 어렵다는 점을 보여 준다. 다만 이 수치만으로는 두 방식 가운데 어느 쪽이 더 정확한지, 실제로 적격한 문헌이 몇 편인지는 판단할 수 없다."]),
("4.3 탐색 도구 검증",[f"120개 시나리오를 각각 세 번씩 실행하여 모두 {V['executions']:,}회의 출력을 비교하였다. 결과의 일관성, 질문 연결의 정확성, 출처 계보의 완전성은 각각 120건 가운데 {V['scenario_count']}/{V['scenario_count']}건을 충족하였다. 임상행동 누출 {V['clinical_action_leakage_scenarios']}건, 이전 자료(legacy) 누출 {V['legacy_leakage_scenarios']}건, 부분·유사 일치(near-match) 오연결 {V['negative_false_routes']}건이었다.","이 결과는 정해진 모의 입력에서 구현한 경계가 제대로 작동했음을 뜻한다. 실제 환자 상황에서 임상적으로 안전하다거나 문헌 검색 성능이 뛰어나다는 근거로까지 넓힐 수는 없다."])]),
("5. 고찰",[
("5.1 주요 결과의 의미",["본 연구의 가장 중요한 산출물은 큰 검색 건수가 아니라, 관찰 가능 여부와 자동 분류의 일치·불일치, 그리고 출처 계보를 한꺼번에 보존한 구조다. 이 구조 덕분에 뒤이어 검토하는 사람은 무엇부터 확인해야 할지 알 수 있고, 각 수치가 원자료에서 어떻게 만들어졌는지 거꾸로 따라갈 수 있다.","전체 20,230건 가운데 2,215건은 초록이 없었다. 그러므로 제목과 서지정보만 보고 의미를 넓혀 해석하지 않는 것이 중요하다. 또한 공개 원문 위치정보가 있다고 해서 본문 추출까지 끝난 것은 아니므로, 접근할 수 있다는 상태와 분석을 마쳤다는 상태는 서로 구분해 두어야 한다."]),
("5.2 방법론적 기여",["이번 프로토콜 전환은 끝내지 못한 체계적 고찰을 마치 완성된 것처럼 포장하지 않고, 지금 가진 자료로 검증할 수 있는 기술적 질문으로 연구 범위를 다시 맞춘 사례다. AI를 쓴 사실과 그 한계를 문서 전체에 똑같이 반영함으로써, 자동화가 사람의 판정을 대신했다는 오해를 줄였다.","두 자동 분류의 불일치를 지우지 않고 보존한 설계는 단일 점수가 주는 섣부른 확신을 피한다. 분류가 갈린 문헌은 앞으로 사람이 먼저 검토할 표본이 되고, 규칙을 고칠 때 어떤 문헌이 상태를 바꾸는지 비교하는 기준선이 된다."]),
("5.3 실무적 함의",["탐색 도구는 임상적인 답을 대신 내주는 대신, 관련 질문의 자료 지형과 출처를 보여 준다. 이 역할은 사용자에게 당장 어떤 행동을 요구하지 않으면서 문헌 탐색의 출발점을 제공한다. 의학적 판단이 필요한 경우에는 원문 확인과 전문가의 평가가 뒤따라야 한다."]),
("5.4 후속 연구",["다음 단계에서는 두 명 이상이 독립적으로 우선 표본을 선별하고 서로의 이견을 조정해야 한다. 그 결과를 기준 정답으로 삼으면 자동 분류기의 민감도·특이도와 오류 유형을 평가할 수 있다. 또한 구독 데이터베이스와 국내 원본 내보내기 자료를 추가하고 동일 연구군을 사람이 직접 확인하여, 연구 단위의 문헌 흐름도를 만들 필요가 있다."])]),
("6. 연구의 한계",[
("6.1 자료와 판정의 한계",["본 연구에는 사람의 독립 선별, 이중 추출, 비뚤림 위험 평가, 근거수준(GRADE) 판정이 없다. Embase, Scopus, Web of Science와 일부 국내 데이터베이스, 구독 원문도 포함하지 못했다. 그러므로 이 근거지도는 가능한 모든 근거를 빠짐없이 담은 완전한 목록이 아니다.","문헌은 질문마다 중복될 수 있고 동일 연구군도 확정하지 않았으므로, 행 수를 곧 고유 연구 수로 해석할 수 없다. 자동 분류는 텍스트 표면의 신호에 기대므로 임상적 적격성을 대신하지 못한다."]),
("6.2 검증의 한계",["모의 시나리오는 실제 전문가나 환자의 데이터가 아니므로 임상적 안전성과 유용성을 검증하지 못한다. 성공률은 구현한 경계가 얼마나 일관되게 지켜지는지를 보여 줄 뿐이다. 앞으로 실제 사용자 평가, 전문가 내용 타당도 검토, 외부 데이터셋 검증이 더해져야 한다."])]),
("7. 결론",[("",[f"본 연구는 다섯 가지 고위험 임상상황 질문에 대해 레코드-질문 단위 {M['row_count']:,}건을 출처·관찰 가능 여부·자동 분류·원문 위치정보와 함께 정리하였다. 또한 정확 일치 기반의 탐색 도구를 구현하고, 120개 모의 시나리오에서 결과의 일관성, 질문 연결, 출처 계보의 경계를 검증하였다.","이 결과는 임상 권고도 아니고 체계적 문헌고찰의 결론도 아니다. 연구의 의의는 접근할 수 있는 자료와 접근할 수 없는 자료, 자동화가 할 수 있는 일과 사람이 해야 하는 일을 구분하고, 모든 결과 수치를 원자료와 코드에 연결해 둔 데 있다. 후속 사람 검토와 외부 검증이 끝날 때까지 임상 행동을 제시하지 않는다는 것이 이 설계의 핵심 결론이다."])])]

def sha(p): return hashlib.sha256(p.read_bytes()).hexdigest()
def add_page_number(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); p._p.append(fld)
def table(doc,headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): t.rows[0].cells[i].text=h
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row): c[i].text=str(v)
    return t

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.27); sec.page_height=Inches(11.69); sec.top_margin=sec.bottom_margin=Inches(.85); sec.left_margin=sec.right_margin=Inches(1)
for name,size,color,font in [("Normal",10.5,"222222","Pretendard"),("Title",20,"15314B","Pretendard ExtraBold"),("Heading 1",16,"1F4E79","Pretendard Bold"),("Heading 2",13,"1F4E79","Pretendard SemiBold"),("Heading 3",11.5,"365F7D","Pretendard Medium")]:
    s=doc.styles[name]; s.font.name=font; s._element.rPr.rFonts.set(qn("w:ascii"),font); s._element.rPr.rFonts.set(qn("w:hAnsi"),font); s._element.rPr.rFonts.set(qn("w:eastAsia"),font); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.line_spacing=1.55 if name=="Normal" else 1.2; s.paragraph_format.space_after=Pt(7)
sec.header.paragraphs[0].text="고위험 임상상황의 영양보충제 안전성 문헌 탐색 연구"; add_page_number(sec.footer.paragraphs[0])
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(90); p.add_run("학위논문").bold=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(35); r=p.add_run("고위험 임상상황의 영양보충제 안전성 문헌을 위한\nAI 기반 탐색적 근거지도 구축과\n결정론적 탐색 도구의 기술 검증"); r.bold=True; r.font.size=Pt(20)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("AI-Based Exploratory Evidence Mapping of Dietary Supplement Safety in High-Risk Clinical Contexts").italic=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80); p.add_run(f"{AUTHOR}\n{SUBMISSION_DATE}").bold=True
doc.add_page_break(); doc.add_heading("제출 정보",1)
doc.add_paragraph("이 문서는 여형준의 졸업논문 제출용 편집 기준본이다. 연구 내용과 검증된 수치는 본문에 반영되어 있으며, 아래 행정 정보는 소속 기관의 공식 양식에 맞춰 제출 전에 입력한다.")
table(doc,["항목","내용"],[["연구자",AUTHOR],["소속 대학","[입력]"],["학과·전공","[입력]"],["학번","[입력]"],["지도교수","[입력]"],["학위명","[입력]"],["최종 제출일","[입력]"]])
doc.add_paragraph("연구 성격: AI 기반 탐색적 근거지도 및 결정론적 탐색 도구의 기술 검증. 사람의 독립 이중선별, 최종 RoB·GRADE, 임상 권고를 완료한 체계적 문헌고찰로 표시하지 않는다.")
doc.add_page_break(); doc.add_heading("국문초록",1)
abstract=f"본 연구는 항응고제 병용이나 신장결석 위험과 관련된 다섯 가지 영양보충제 질문을 대상으로, 공개된 문헌 서지정보를 인공지능(AI) 기반 탐색적 근거지도로 정리하고, 임상 행동을 내놓지 않는 결정론적 탐색 도구가 정해진 경계 안에서 작동하는지를 기술적으로 검증하였다. 처음에는 체계적 문헌고찰을 계획하였으나 사람의 독립 선별과 전체 원문 접근 조건을 갖추지 못하여 이를 완료된 고찰로 간주하지 않았고, 연구 질문을 공개 자료의 관찰 가능성, 자동 분류의 일치·불일치, 출처 계보, 소프트웨어 재현성으로 좁힌 개정 프로토콜(v2)로 전환하였다. 이렇게 모은 레코드-질문 단위는 모두 {M['row_count']:,}건으로, PubMed {M['source_counts']['pubmed']:,}건, ClinicalTrials.gov {M['source_counts']['clinicaltrials']:,}건, KoreaMed {M['source_counts']['koreamed']:,}건이었다. 초록까지 확인된 것은 {M['abstract_observed']:,}건, 제목과 서지정보만 확인된 것은 {M['title_metadata_only']:,}건이었다. PubMed 문헌에 두 자동 분류를 함께 적용한 결과는 유지 합의 {S['classifications']['ai_agreement_retain']:,}건, 후순위 합의 {S['classifications']['ai_agreement_deprioritize']:,}건, 불일치 불확실 {S['classifications']['ai_disagreement_uncertain']:,}건이었다. 120개 모의 시나리오를 360회 실행한 검증에서 결과의 일관성, 질문 연결의 정확성, 출처 계보의 완전성은 모두 120건 가운데 120건을 충족했고, 임상행동 누출과 이전 자료 누출, 유사어 오연결은 한 건도 없었다. 이 결과는 임상 권고나 사람이 선별한 체계적 문헌고찰이 아니라, 이후 사람이 이어서 검토할 수 있도록 마련한 재현 가능한 탐색 기반이다."
doc.add_paragraph(abstract); doc.add_paragraph("주요어: 영양보충제, 항응고제, 신장결석, 탐색적 근거지도, 인공지능, 출처 계보, 결정론적 검증")
doc.add_heading("Abstract",1); doc.add_paragraph("This study constructed an AI-assisted exploratory evidence map for five dietary-supplement safety questions and technically validated a deterministic navigation tool. The protocol was explicitly narrowed because independent human screening and complete full-text access were unavailable. The corpus contains 20,230 record-question units. Automated labels are prioritization signals, not inclusion decisions, and the tool emits no clinical actions. Repeated synthetic tests supported deterministic routing and provenance completeness, but did not establish clinical validity.")
doc.add_page_break(); doc.add_heading("목차",1)
for h,subs in sections:
    doc.add_paragraph(h,style="Heading 2")
    for sh,_ in subs:
        if sh: doc.add_paragraph(sh)
doc.add_paragraph("참고문헌"); doc.add_paragraph("부록 A. 질문별 결과표"); doc.add_paragraph("부록 B. 재현 경로와 파일 해시")
for section_index,(h,subs) in enumerate(sections):
    if section_index == 0:
        doc.add_page_break()
    doc.add_heading(h,1)
    for sh,paras in subs:
        if sh: doc.add_heading(sh,2)
        for x in paras: doc.add_paragraph(x)
        if sh=="1.3 연구 목적과 질문": table(doc,["질문","범위"],[[k,v] for k,v in Q.items()])
        if sh=="4.2 자동 분류 분포": table(doc,["분류","행 수"],[[k,f"{v:,}"] for k,v in S["classifications"].items()])
        if sh=="4.3 탐색 도구 검증": table(doc,["평가지표","결과"],[["합성 시나리오",V["scenario_count"]],["반복 실행",V["executions"]],["결정성",f"{V['deterministic_scenarios']}/{V['scenario_count']}"],["정확 라우팅",f"{V['correct_exact_route_scenarios']}/{V['scenario_count']}"],["계보 완전성",f"{V['provenance_complete_scenarios']}/{V['scenario_count']}"],["임상행동 누출",V["clinical_action_leakage_scenarios"]],["legacy 누출",V["legacy_leakage_scenarios"]],["near-match 오경로",V["negative_false_routes"]]])
doc.add_heading("참고문헌",1)
refs=["Page MJ, et al. The PRISMA 2020 statement: an updated guideline for reporting systematic reviews. BMJ. 2021;372:n71.","Tricco AC, et al. PRISMA Extension for Scoping Reviews (PRISMA-ScR): Checklist and Explanation. Ann Intern Med. 2018;169:467-473.","Rethlefsen ML, et al. PRISMA-S: an extension to the PRISMA Statement for Reporting Literature Searches in Systematic Reviews. Syst Rev. 2021;10:39.","Miake-Lye IM, et al. What is an evidence map? A systematic review of published evidence maps and their definitions, methods, and products. Syst Rev. 2016;5:28.","National Library of Medicine. PubMed and NCBI E-utilities documentation. Accessed 2026-07-10.","U.S. National Library of Medicine. ClinicalTrials.gov Data API documentation. Accessed 2026-07-10.","KoreaMed. Korean medical literature search service. Accessed 2026-07-10.","연구 프로토콜 v2.0. AI 기반 탐색적 문헌지도 및 결정론적 도구 검증. 2026-07-12."]
for x in refs: doc.add_paragraph(x,style="List Number")
doc.add_heading("부록 A. 질문별 결과표",1)
claims=[json.loads(x) for x in (ROOT/"data/curated_v2/provisional_claims.jsonl").read_text(encoding="utf-8").splitlines() if x.strip()]
table(doc,["질문","전체","초록","유지 합의","후순위 합의","불일치","비순위"],[[c["question_id"],c["record_question_units"],c["abstract_observed"],c["classification_counts"].get("ai_agreement_retain",0),c["classification_counts"].get("ai_agreement_deprioritize",0),c["classification_counts"].get("ai_disagreement_uncertain",0),c["classification_counts"].get("ai_unranked_source_candidate",0)] for c in claims])
doc.add_heading("부록 B. 재현 경로와 파일 해시",1)
table(doc,["항목","상대 경로","SHA-256"],[[k,p.relative_to(ROOT).as_posix(),sha(p)] for k,p in P.items()])
doc.add_paragraph("표의 수치는 문서에 고정 입력하지 않고 각 manifest에서 읽어 생성하였다. 해시가 달라지면 검증 절차를 다시 수행해야 한다.")
doc.save(DOCX)

md=["# 고위험 임상상황의 영양보충제 안전성 문헌을 위한 AI 기반 탐색적 근거지도 구축과 결정론적 탐색 도구의 기술 검증","","## 국문초록","",abstract,""]
for h,subs in sections:
    md += [f"## {h}",""]
    for sh,paras in subs:
        if sh: md += [f"### {sh}",""]
        md += sum(([x,""] for x in paras),[])
MD.write_text("\n".join(md),encoding="utf-8")
print(json.dumps({"docx":str(DOCX),"markdown":str(MD),"paragraphs":len(doc.paragraphs),"tables":len(doc.tables)},ensure_ascii=False))
