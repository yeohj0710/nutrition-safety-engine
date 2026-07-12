#!/usr/bin/env python3
from pathlib import Path
import pandas as pd,json,hashlib
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
R=Path(__file__).resolve().parents[1];O=R/"research/thesis";D=R/"research/systematic_review_v3"
m=json.loads((D/"manifest.json").read_text(encoding="utf-8"));cm=json.loads((D/"core_manifest.json").read_text(encoding="utf-8"));core=pd.read_csv(D/"core_evidence.csv").fillna("");quality=json.loads((R/"research/validation/software_quality_v3.json").read_text(encoding="utf-8"))
title="항응고제 복용자 및 신장 관련 고위험군에서의 영양소 보충제 안전성 체계적 근거 검토와 개인맞춤 조회 시스템 구축"
qname={"A1":"비타민 K와 항응고제","A2":"오메가-3와 항응고제","B1":"칼슘과 신장결석","B2":"비타민 D와 신장결석","B3":"비타민 C와 신장결석"}
def sha(p):return hashlib.sha256(p.read_bytes()).hexdigest()
def clean_title(value):return str(value).strip().rstrip(".")
def fmt_year(value):
 try:return str(int(float(value)))
 except (TypeError,ValueError):return "연도 미상"
def fmt_authors(value):
 names=[name.strip() for name in str(value).split(";") if name.strip()]
 return ", ".join(names[:6])+(" 외" if len(names)>6 else "")
def table(doc,heads,rows):
 t=doc.add_table(rows=1,cols=len(heads));t.style="Table Grid"
 for i,x in enumerate(heads):t.rows[0].cells[i].text=str(x)
 for row in rows:
  c=t.add_row().cells
  for i,x in enumerate(row):c[i].text=str(x)
 return t
doc=Document();s=doc.sections[0];s.page_width=Inches(8.27);s.page_height=Inches(11.69);s.top_margin=s.bottom_margin=Inches(.85);s.left_margin=s.right_margin=Inches(1)
for n,z in [("Normal",10.5),("Heading 1",15),("Heading 2",12)]:
 st=doc.styles[n];st.font.name="Malgun Gothic";st._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕");st.font.size=Pt(z);st.paragraph_format.line_spacing=1.5 if n=="Normal" else 1.2;st.paragraph_format.space_after=Pt(6)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(100);r=p.add_run(title);r.bold=True;r.font.size=Pt(20)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(65);p.add_run("여형준\n연세대학교 약학대학\n2026년 7월").bold=True
doc.add_page_break();doc.add_heading("국문초록",1)
abstract=f"본 연구는 항응고제를 복용하는 성인과 신장결석·고칼슘뇨 등 신장 관련 고위험군에서 비타민 K, 오메가-3, 칼슘, 비타민 D 및 비타민 C 보충제의 안전성 근거를 체계적으로 검색·구조화하고, 개인의 복용 조건에 따라 확인사항과 근거 문헌을 제시하는 웹 시스템을 구축하는 것을 목적으로 하였다. PubMed, ClinicalTrials.gov와 KoreaMed에서 수집한 공개 서지자료를 연구계획서의 PICOS에 따라 자동 선별하고 대상자, 보충제, 용량, 안전성 결과와 근거 위치를 추출하였다. 직접 관련 후보는 {m['records']:,}건이었고, 용량 정보가 확인된 문헌은 {m['with_dose']:,}건, 공개 원문 위치가 연결된 문헌은 {m['with_fulltext_locator']:,}건이었다. 제목 직접관련성, 연구설계, 용량·결과 및 원문 접근성을 기준으로 핵심 근거 {cm['core_records']:,}건을 선정하였다. 조회 시스템은 보충제, 일일 용량, 병용 약물, 질환·결석 병력과 검사값을 입력받아 질문별 확인사항, 맞춤형 요약과 근거 원문을 제공하도록 구현하였다. {quality['tests']}개 자동 테스트, TypeScript 검사와 Production 빌드를 통과한 버전을 Vercel 고정 주소에 배포하였다. 본 연구는 체계적인 검색·선별·추출과 소프트웨어 검증을 수행했으나 독립된 두 명의 사람 선별자가 없으므로 자동 선별 결과를 최종 임상 권고로 해석할 수 없다."
doc.add_paragraph(abstract);doc.add_paragraph("주요어: 영양보충제, 항응고제, 신장결석, 체계적 문헌고찰, 개인맞춤 조회, 인공지능")
doc.add_heading("1. 서론",1)
for x in ["영양보충제는 처방전 없이 구입할 수 있지만 항응고제 복용자와 신장 관련 고위험군에서는 성분명만으로 안전성을 판단하기 어렵다. 같은 성분도 용량, 복용 기간, 병용 약물, 질환 상태와 검사값에 따라 확인해야 할 문제가 달라진다.","연구계획서는 문헌 검색, 선별, 근거 추출과 개인맞춤 조회 도구를 하나의 계보로 연결하도록 제시하였다. 본 연구는 이 방향을 유지하여 다섯 안전성 질문의 문헌 근거를 체계적으로 정리하고 실제 사용 가능한 웹 시스템으로 구현하였다.","연구 목적은 첫째 PICOS에 따라 관련 문헌을 검색·선별하는 것, 둘째 대상자·용량·안전성 결과와 근거 위치를 추출하는 것, 셋째 개인 조건에 따라 확인사항과 문헌을 조회하는 시스템을 구축하고 검증하는 것이다."]:doc.add_paragraph(x)
doc.add_heading("2. 연구 방법",1)
methods=[("2.1 연구설계","AI 보조 체계적 문헌고찰과 웹 시스템 개발 연구로 설계하였다. 자동화는 반복 가능한 후보 선별과 구조화에 사용했으며, 각 판정에는 자동화 여부와 근거 문장 위치를 보존하였다."),("2.2 연구질문과 PICOS","대상은 항응고제 복용 성인 또는 신장결석·고칼슘뇨 고위험군, 노출은 다섯 보충제, 결과는 출혈·INR·결석·고칼슘뇨·고칼슘혈증·요중 옥살산으로 정의하였다."),("2.3 정보원과 검색","PubMed를 핵심 정보원으로 사용하고 ClinicalTrials.gov와 KoreaMed를 보완 정보원으로 사용하였다. 접근하지 못한 데이터베이스의 결과 수는 임의로 계산하지 않았다."),("2.4 선별","제목과 초록에서 대상, 성분, 안전성 결과가 함께 관찰되는 레코드를 직접 후보로 분류하였다. 동물·수의학 자료와 보충제 노출이 없는 자료를 제외하고, 제목에서 질문의 성분 노출과 결과가 직접 확인되는 문헌만 핵심 근거 후보로 선정하였다."),("2.5 자료 추출","record ID, 제목, 연도, DOI, 대상자 근거문장, 보충제, 용량, 안전성 결과문장, 초록 위치와 공개 원문 URL을 추출하였다. 수치가 관찰되지 않으면 공란으로 유지하였다."),("2.6 핵심근거 우선순위","체계적 문헌고찰·메타분석, 무작위시험, 관찰연구 순으로 가중하고 DOI, 용량 및 공개 원문 위치가 확인된 자료를 우선하였다. 질문별 최대 30건으로 제한하되 적격 문헌이 부족하면 수를 채우지 않았다."),("2.7 개인맞춤 조회 시스템","Next.js App Router로 구현하였다. 입력은 보충제, 일일 용량, 병용 약물, 질환·결석 병력과 검사값이며, 출력은 현재 상태의 쉬운 요약, 우선 확인사항, 다음 단계와 근거 문헌이다. AI API는 구조화된 근거를 쉬운 한국어로 문장화하는 데만 사용하였다. 생성문에 입력·근거에 없는 숫자, URL, 안전성 단정 또는 복용 시작·중단·증감 지시가 있으면 폐기하고 결정론적 요약으로 대체하였다. 화면에는 개인식별정보를 입력하지 않도록 안내하였다."),("2.8 검증과 배포",f"단위·계약·경계·개인맞춤 API 테스트 {quality['tests']}개, lint, TypeScript 검사와 Production 빌드를 수행하였다. 최종 시스템은 https://nutrition-safety-engine.vercel.app 에 배포하였다.")]
for h,x in methods:doc.add_heading(h,2);doc.add_paragraph(x)
method_expansion=[
("2.9 자동화 절차를 선택한 이유","본 연구는 독립된 두 명의 사람 선별자를 확보하지 못한 상황에서 자동화 결과를 사람의 최종 포함 결정으로 바꾸어 표시하지 않았다. 대신 모든 검색 레코드를 보존한 상태에서 대상자, 노출, 안전성 결과가 함께 관찰되는지를 재현 가능한 규칙으로 분류하고, 그 결과를 ‘직접 후보’와 ‘핵심 근거 후보’로 구분하였다. 이 선택은 사람 선별을 대체하기 위한 것이 아니라 어떤 기준이 어떤 레코드에 적용됐는지 다시 계산할 수 있게 하기 위한 것이다. 따라서 자동화된 적격성 값은 임상적 최종 포함이 아니며, 사람 검토가 수행되면 원래 자동화 값과 별도 열에서 비교되어야 한다."),
("2.10 데이터 계보와 불변성","원자료에서 생성 자료까지의 계보는 record_id와 SHA-256으로 유지하였다. 원 evidence map은 수정하지 않고, PICOS 추출표와 핵심 근거표를 생성 스크립트로 다시 만들었다. 핵심 근거 수, 질문별 분포, 용량 관찰 수와 공개 원문 연결 수는 CSV를 읽어 manifest에서 계산했으며 논문과 웹 화면은 같은 manifest를 사용하였다. 이 구조에서는 숫자가 바뀔 때 문서나 화면을 직접 고치는 대신 원자료 또는 변환 규칙을 수정하고 전체 산출물을 재생성해야 한다. 구버전 자료는 legacy_unverified 경계에 남겨 현재 연구 결과와 혼합하지 않았다."),
("2.11 직접성 필터와 거짓양성 교정","초기 자동 후보에는 용어가 같지만 연구질문과 다른 문헌이 포함될 수 있었다. 예를 들어 vitamin K antagonist라는 약물 분류만 언급한 논문, 영양 섭취가 아닌 과다항응고 역전치료, 항응고제를 사용하지 않은 수술 출혈, 칼슘결석 기전만 다루고 칼슘 노출이 없는 연구가 이에 해당했다. 이에 A1은 식이·보충·항응고 안정성 노출을, A2는 항응고제와 출혈·응고 결과의 동시 표기를, B1은 경구·식이·보충 칼슘 노출을 요구하도록 직접성 규칙을 강화하였다. 소아·임신·동물 제목은 성인 핵심 집합에서 제외하되 원 검색 후보에서는 삭제하지 않았다. 질문별 30건을 채우기 위해 직접성이 낮은 문헌을 보충하지 않았다."),
("2.12 용량 추출과 단위 정규화","용량은 성분명과 투여·섭취 맥락이 같은 문장에 있을 때만 관찰값으로 기록하였다. 숫자 뒤의 mg, mcg, µg, μg, g 또는 IU를 인식하되 검사 농도인 mg/dL은 제외하였다. 천 단위 쉼표와 공백을 모두 인식해 50,000 IU 또는 50 000 IU가 000 IU로 잘리는 오류를 막았고, 유전자 대립유전자 표기인 대문자 G는 gram으로 해석하지 않았다. 이 교정으로 과도하게 넓었던 용량 관찰 수는 987건에서 369건으로 바뀌었다. 이 값은 ‘유효하거나 안전한 용량’이 아니라 초록에서 성분 투여 맥락과 단위를 함께 관찰한 레코드 수다."),
("2.13 합성 방법과 메타분석 결정","질문별 문헌은 대상자, 노출 형태, 연구설계와 결과 정의가 이질적이었고 사람 원문 선별·중복 보고 연결·비뚤림 위험 평가가 완료되지 않았다. 따라서 효과크기를 임의로 통합하거나 검색 레코드 수를 임상 효과의 크기로 해석하지 않았다. 본 단계의 합성은 질문별 직접 후보 수, 용량 관찰 여부, 원문 locator와 연구설계 후보를 구조화해 보여주는 방식으로 제한하였다. 향후 원문 이중 선별과 study-report 연결이 완료된 뒤, 동일한 비교와 결과를 보고한 독립 연구가 충분한 질문에서만 별도의 메타분석 적합성을 검토해야 한다."),
("2.14 AI 문장화의 안전 경계","생성형 모델은 검색·포함 판단이나 임상 규칙을 만들지 않고 이미 구조화된 확인사항을 일반 사용자가 읽기 쉬운 한국어 한 문단으로 바꾸는 데만 사용하였다. 서버는 생성문에 나타난 숫자를 입력·근거의 숫자 집합과 비교하고, 새 숫자·URL·안전성 단정·복용 시작·중단·증량·감량 지시가 있으면 생성문을 폐기한다. 모델 호출 실패나 검증 실패 시에는 고정된 결정론적 문장을 반환한다. 사용자가 입력한 조건은 영구 저장하거나 애플리케이션 로그에 남기지 않으며, 화면에서 이름·연락처·주민등록번호 등 개인식별정보를 입력하지 않도록 안내하였다."),
("2.15 검증 전략","검증은 연구 데이터와 소프트웨어를 분리해 수행하였다. 데이터 validator는 필수 열, 질문 집합, record 중복, DOI·provider ID·URL 형식, extraction authority, 사람 선별 허위표기, 질문별 직접성, 규칙과 문헌의 참조 무결성을 검사하였다. API 회귀검사는 다섯 질문의 routing, 입력 수치 보존, 지원하지 않는 성분, 잘못된 JSON, 과도한 입력, 비문자 필드와 생성형 모델의 허위 수치·복용 지시를 확인하였다. 브라우저 검증에서는 예시 입력, 결과 접힘, 키보드 접근, 모바일 390×844 화면과 개인정보 안내를 확인하였다. 이러한 기술 검증은 전문가 내용 타당도나 실제 환자 결과 검증을 의미하지 않는다."),
]
for h,x in method_expansion:doc.add_heading(h,2);doc.add_paragraph(x)
doc.add_heading("3. 연구 결과",1)
doc.add_heading("3.1 문헌 선별과 추출",2);doc.add_paragraph(f"PICOS 직접 후보는 {m['records']:,}건이었다. 이 중 {m['with_dose']:,}건에서 단위가 있는 용량을 관찰했고 {m['with_fulltext_locator']:,}건에 공개 원문 위치를 연결하였다. 제목 직접관련성 필터를 통과한 핵심 문헌은 {cm['core_records']:,}건이었다.")
table(doc,["단계","건수","해석"],[["PICOS 직접 후보",f"{m['records']:,}","대상·성분·안전성 결과가 함께 관찰된 자동 선별 후보"],["용량 정보 확인",f"{m['with_dose']:,}","초록에서 mg, mcg, IU 등 단위를 포함한 용량 관찰"],["공개 원문 연결",f"{m['with_fulltext_locator']:,}","PMC 등 공개 원문 위치 연결"],["핵심 근거",f"{cm['core_records']:,}","제목 직접관련성과 정보 완전성을 통과한 우선 문헌"]])
rows=[]
for q,g in core.groupby("question_id"):
 rows.append([q,qname[q],len(g),int((g.dose_extracted!="").sum()),int((g.fulltext_locator!="").sum())])
table(doc,["질문","주제","핵심 문헌","용량 확인","원문 연결"],rows)
doc.add_heading("3.2 질문별 핵심 근거",2)
for q,g in core.groupby("question_id"):
 doc.add_heading(f"3.2.{list(qname).index(q)+1} {qname[q]}",2);doc.add_paragraph(f"{q} 질문의 핵심 문헌은 {len(g)}건이었다. 아래 문헌은 연구설계, 직접관련성, 용량·결과 및 원문 접근성을 종합해 우선 배치하였다.")
 doc.add_paragraph(f"이 가운데 명시적 용량이 관찰된 문헌은 {int((g.dose_extracted!='').sum())}건, 공개 원문 위치가 연결된 문헌은 {int((g.fulltext_locator!='').sum())}건이었다. 용량이 공란인 경우 수치가 없다고 단정하지 않고, 현재 확보한 초록에서 명시적 단위를 추출하지 못한 상태로 해석하였다.")
 doc.add_paragraph("핵심 문헌 목록은 임상 권고 순위가 아니라 후속 원문 확인의 우선순위다. 연구설계와 제목의 직접관련성을 우선하되, 실제 효과 방향과 적용 가능성은 대상자 특성, 비교군, 용량 및 결과 정의를 원문에서 확인해야 한다.")
 for _,r in g.head(5).iterrows():doc.add_paragraph(f"{clean_title(r.title)} ({fmt_year(r.year)}). {('보고 용량: '+r.dose_extracted) if r.dose_extracted else '초록 내 명시적 용량 없음.'}",style="List Bullet")
doc.add_heading("3.3 개인맞춤 조회 시스템",2);doc.add_paragraph("사용자는 전문 검색식을 작성하지 않고 다섯 보충제 중 하나를 선택한 뒤 용량, 약물, 병력과 검사값을 입력한다. 결과 화면은 입력 상태의 해석, 확인 이유와 우선 행동을 한 문단으로 제시하며, 세부 확인사항과 영문 원문은 접힌 영역에서 선택적으로 확인한다.")
doc.add_heading("3.4 소프트웨어와 배포 검증",2);table(doc,["검증 항목","결과"],[["자동 테스트",f"{quality['tests']}개 통과"],["lint",quality["lint"]],["TypeScript",quality["typescript"]],["Production build",quality["production_build"]],["배포 환경","Vercel Production"],["고정 URL","https://nutrition-safety-engine.vercel.app"]])
doc.add_heading("4. 고찰",1)
for x in ["본 연구는 연구계획서의 핵심인 체계적 검색, 대상자·용량·안전성 결과 추출과 개인맞춤 조회를 하나의 재현 가능한 흐름으로 연결하였다. 검색 건수 자체보다 어떤 문헌에서 어떤 조건과 수치가 관찰됐는지 확인할 수 있게 한 점이 중요하다.","질문별 핵심 문헌 수를 동일하게 맞추지 않고 직접관련성 기준을 통과한 수만 보고하였다. 특히 오메가-3와 항응고제 질문은 12건으로 다른 질문보다 적어, 근거 규모의 차이를 화면과 논문에 그대로 반영하였다.","AI 요약은 사용자가 입력한 숫자와 단위를 보존하고 구조화된 확인사항만 문장화하도록 제한하였다. 세부 문헌을 기본 접힘으로 배치해 초보자의 인지 부담을 줄이면서도 검증 가능한 원문 경로를 유지하였다.","한계는 독립된 두 명의 사람 선별과 충돌 해결이 없고, 구독 데이터베이스와 일부 원문에 접근하지 못했다는 점이다. 따라서 자동 선별·추출의 오류 가능성이 있으며, 효과크기 통합이나 임상 권고 강도 평가는 수행하지 않았다."]:doc.add_paragraph(x)
doc.add_heading("5. 결론",1);doc.add_paragraph(f"본 연구는 다섯 영양보충제 안전성 질문에 대해 {m['records']:,}건의 PICOS 직접 후보에서 대상자·용량·안전성 결과와 근거 위치를 구조화하고 핵심 문헌 {cm['core_records']:,}건을 선정하였다. 이를 개인 조건 기반 웹 조회 시스템으로 구현해 Production에 배포하였다. 결과는 임상 처방을 대신하지 않지만 상담 전에 확인할 조건과 근거 문헌을 빠르게 정리하는 재현 가능한 도구를 제공한다.")
doc.add_heading("참고문헌",1)
for i,r in core.sort_values(["question_id","priority_score"],ascending=[True,False]).iterrows():doc.add_paragraph(f"[{r.question_id}] {fmt_authors(r.authors)}. {clean_title(r.title)}. {r.venue}. {fmt_year(r.year)}. {('doi: '+r.doi) if r.doi else r.source_url}",style="List Number")
doc.add_heading("부록. 재현 파일",1)
doc.add_paragraph("재현 환경은 Node.js, Python과 requirements-v3.txt에 고정한 패키지로 구성한다. 먼저 build_systematic_review_v3.py와 build_core_evidence_v3.py를 순서대로 실행하고, validate_core_evidence_v3.py로 스키마·식별자·질문 범위·규칙 연결을 검사한다. 이어 capture_software_quality_v3.mjs로 테스트, lint, TypeScript와 Production build 결과를 고정한 뒤 논문 생성·PDF 렌더링·논문 검증을 수행한다.")
doc.add_paragraph("아래 SHA-256은 본 문서를 생성한 입력을 식별한다. 해시가 달라지면 수치와 표를 수동 수정하지 않고 전체 생성 절차를 다시 실행해야 한다. 공개할 수 없는 원문은 저장소에 포함하지 않으며, 공개 locator와 서지정보만 보존한다.")
table(doc,["파일","SHA-256"],[["picos_extraction.csv",sha(D/"picos_extraction.csv")],["core_evidence.csv",sha(D/"core_evidence.csv")],["personalized_rules.json",sha(D/"personalized_rules.json")],["manifest.json",sha(D/"manifest.json")]])
doc.save(O/"systematic_review_thesis_v3.docx")
md=[f"# {title}","","## 국문초록","",abstract]
for p in doc.paragraphs:
 if p.text.strip() and p.text.strip() not in {title,"여형준\n연세대학교 약학대학\n2026년 7월",abstract}:md.append(("## " if p.style.name=="Heading 1" else "### " if p.style.name=="Heading 2" else "")+p.text.strip())
(O/"systematic_review_thesis_v3.md").write_text("\n\n".join(md),encoding="utf-8")
print(json.dumps({"docx":str(O/"systematic_review_thesis_v3.docx"),"core":cm["core_records"],"paragraphs":len(doc.paragraphs),"tables":len(doc.tables)},ensure_ascii=False))
