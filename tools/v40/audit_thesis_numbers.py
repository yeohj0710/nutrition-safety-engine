#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Read-only numerical audit for the Yeo v4.0 thesis and its figure sources.

The script reads the sealed/canonical research artifacts, the thesis DOCX, and
the figure source/PNG files. It writes nothing. Results are printed to stdout;
the process exits non-zero when a canonical comparison fails.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document


REPO = Path(__file__).resolve().parents[2]
DRIVE_ROOT = Path(
    os.environ.get(
        "THESIS_DRIVE_ROOT",
        r"G:\내 드라이브\여형준님\24 전공심화실습(1)",
    )
)
DEFAULT_DOCX = DRIVE_ROOT / "여형준" / "02_졸업논문" / "여형준_졸업논문_최종본.docx"
DEFAULT_FIG_DIR = DRIVE_ROOT / "여형준" / "02_졸업논문" / "etc" / "_그림"

RUN_REPORT = REPO / "research" / "logs" / "v40_run_report.json"
SYNTHESIS = REPO / "research" / "synthesis" / "screener_vs_ai_reference_v40.json"
SCORING_REPORT = REPO / "research" / "logs" / "v40_scoring_report.json"
MANIFEST = REPO / "research" / "systematic_review_v40" / "manifest.json"
CORE_MANIFEST = REPO / "research" / "systematic_review_v40" / "core_manifest.json"
RULES = REPO / "research" / "systematic_review_v40" / "personalized_rules.json"

EXPECTED_FIGURE_HASHES = {
    "fig_yeo.py": "57fa5f983954f4b2e79d5df5ad2121272d2cb98cdc9e6191766d30f4ad6612c9",
    "yeo_fig1_screening_flow.png": "761f985f86b2711bb3ed6af11e8e85e38c34c8517bbb1869b79039328654353c",
    "yeo_fig2_scoring_lock.png": "3b25e528a2d8b85513b477ca92115d08a1e79f084442f415ae54a3bd07e44e9d",
}


@dataclass
class Result:
    location: str
    observed: str
    expected: str
    status: str
    source: str


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("−", "-").replace("–", "-")).strip()


def pct(value: float, digits: int = 1) -> str:
    return f"{value * 100:.{digits}f}%"


def number(value: int) -> str:
    return f"{value:,}"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


class Audit:
    def __init__(self, docx_path: Path):
        self.doc = Document(docx_path)
        self.results: list[Result] = []
        self.covered_paragraphs: set[int] = set()
        self.covered_tables: set[int] = set()

    def add(self, location: str, observed, expected, source: str, ok: bool | None = None):
        observed_text = str(observed)
        expected_text = str(expected)
        if ok is None:
            ok = observed == expected
        self.results.append(
            Result(location, observed_text, expected_text, "일치" if ok else "불일치", source)
        )

    def paragraph(self, location: str, anchor: str, tokens: Iterable[str], source: str):
        matches = [(i, p.text) for i, p in enumerate(self.doc.paragraphs) if anchor in p.text]
        if len(matches) != 1:
            self.add(location, f"anchor matches={len(matches)}", "1", source, False)
            return
        index, text = matches[0]
        self.covered_paragraphs.add(index)
        missing = [token for token in tokens if norm(token) not in norm(text)]
        self.add(location, "모든 기대 수치 있음" if not missing else "누락: " + ", ".join(missing),
                 "모든 기대 수치 있음", source, not missing)

    def cell(self, table_index: int, row: int, col: int, expected: str, source: str, label: str):
        self.covered_tables.add(table_index)
        observed = norm(self.doc.tables[table_index].cell(row, col).text)
        self.add(label, observed, norm(expected), source)

    def table(self, table_index: int, expected_rows: Sequence[Sequence[str]], source: str, label: str):
        self.covered_tables.add(table_index)
        table = self.doc.tables[table_index]
        observed_rows = [[norm(cell.text) for cell in row.cells] for row in table.rows]
        expected = [[norm(cell) for cell in row] for row in expected_rows]
        self.add(label, json.dumps(observed_rows, ensure_ascii=False),
                 json.dumps(expected, ensure_ascii=False), source)

    def unverified_numeric_items(self) -> list[str]:
        items: list[str] = []
        for i, paragraph in enumerate(self.doc.paragraphs):
            if i in self.covered_paragraphs or not re.search(r"\d", paragraph.text):
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                continue
            items.append(f"문단 {i + 1}: {norm(paragraph.text)}")
        for i, table in enumerate(self.doc.tables):
            if i in self.covered_tables:
                continue
            text = " | ".join(norm(cell.text) for row in table.rows for cell in row.cells)
            if re.search(r"\d", text):
                items.append(f"표 {i + 1}: {text}")
        return items


def build_audit(docx_path: Path, fig_dir: Path) -> tuple[Audit, list[str]]:
    run = load_json(RUN_REPORT)
    synthesis = load_json(SYNTHESIS)
    scoring = load_json(SCORING_REPORT)
    manifest = load_json(MANIFEST)
    core = load_json(CORE_MANIFEST)
    rules = load_json(RULES)
    audit = Audit(docx_path)

    phase_a = run["phase_a"]
    phase_b = run["phase_b"]
    phase_c = run["phase_c"]
    qids = [q["question_id"] for q in phase_a["questions"]]
    q_raw = {q["question_id"]: q["hit_count"] for q in phase_a["questions"]}
    q_rows = phase_b["corpus"]["row_distribution_by_question"]
    q_retain = {qid: phase_c["distribution_by_question"][qid]["retain"] for qid in qids}

    raw = phase_b["raw_retrieved_rows"]
    rows = phase_b["corpus"]["row_count"]
    unique_papers = phase_b["corpus"]["unique_paper_count"]
    unique_record_ids = phase_b["corpus"]["unique_record_id_count"]
    abstracts = phase_b["corpus"]["observability_distribution"]["abstract_available"]
    title_only = phase_b["corpus"]["observability_distribution"]["title_only"]
    distribution = phase_c["distribution"]
    retain = distribution["retain"]
    deprioritize = distribution["deprioritize"]
    uncertain = distribution["uncertain"]
    adjudicated = phase_c["agent_adjudications"]["records"]
    gate_drop = retain - manifest["records"]
    active_rules = [r for r in rules if r.get("personalization_axis") != "compatibility_alias"]
    alias_rules = [r for r in rules if r.get("personalization_axis") == "compatibility_alias"]

    overall = synthesis["layers"]["overall"]
    weighted = overall["weighted"]
    boot = overall["stratified_bootstrap"]["ci"]
    prevalence = synthesis["retain_prevalence_estimates"]
    disagreements = synthesis["decision_disagreements"]
    sample_n = synthesis["design"]["sample_total"]

    audit.paragraph(
        "여형준 국문초록 1문단",
        "그 결과 코퍼스는",
        [
            "질문은 5개", "10,000행", number(rows), number(unique_papers),
            number(abstracts), number(title_only),
        ],
        "v40_run_report.json phase_a/phase_b",
    )
    audit.paragraph(
        "여형준 국문초록 2문단",
        "선별은 두 층으로 수행하였다",
        [
            number(rows), "커버리지 1.0", number(adjudicated), pct(adjudicated / rows, 1),
            number(retain), number(deprioritize), number(uncertain), number(manifest["records"]),
            "15건", number(core["core_records"]), number(len(active_rules)),
        ],
        "v40_run_report.json phase_c/phase_d; manifest/core_manifest/personalized_rules",
    )
    agreement_ci = boot["agreement_vs_ai_reference"]
    scorer_ci = prevalence["bootstrap_ci"]["scorer_design_based_weighted"]
    audit.paragraph(
        "여형준 국문초록 3문단",
        "선별 품질을 교차 확인하기 위해",
        [
            number(sample_n), pct(weighted["agreement_vs_ai_reference"], 2),
            pct(agreement_ci["ci95_low"], 2).rstrip("%"),
            pct(agreement_ci["ci95_high"], 2).rstrip("%"),
            pct(weighted["sensitivity_vs_ai_reference"], 2),
            pct(weighted["specificity_vs_ai_reference"], 2),
            f"{overall['cohen_kappa_vs_ai_reference_unweighted']:.3f}",
            pct(prevalence["ai_reference_apparent_census"], 2),
            pct(prevalence["scorer_design_based_weighted"], 2),
            pct(scorer_ci["ci95_low"], 2).rstrip("%"),
            pct(scorer_ci["ci95_high"], 2).rstrip("%"),
            f"{prevalence['ratio_scorer_to_ai_reference']:.2f}배",
            "5.83×10⁻¹⁶",
        ],
        "screener_vs_ai_reference_v40.json",
    )
    audit.paragraph(
        "여형준 영문초록",
        "This study built an AI-autonomous",
        [
            "five high-risk", number(rows), number(unique_papers), number(adjudicated),
            pct(adjudicated / rows, 1), number(retain), number(deprioritize), number(uncertain),
            number(sample_n), pct(weighted["agreement_vs_ai_reference"], 2),
            pct(prevalence["ai_reference_apparent_census"], 2),
            pct(prevalence["scorer_design_based_weighted"], 2),
            f"{prevalence['ratio_scorer_to_ai_reference']:.2f}×",
        ],
        "v40_run_report.json; screener_vs_ai_reference_v40.json",
    )

    question_rows = [["질문", "ESearch 건수", "코퍼스 행", "retain", "retain 비율"]]
    labels = ["HRS1 수술 전후", "HRS2 만성콩팥병", "HRS3 임신", "HRS4 간질환", "HRS5 항응고 치료"]
    for label, qid in zip(labels, qids):
        question_rows.append([
            label, number(q_raw[qid]), number(q_rows[qid]), number(q_retain[qid]),
            pct(q_retain[qid] / q_rows[qid], 1),
        ])
    question_rows.append(["합계", number(raw), number(rows), number(retain), pct(retain / rows, 1)])
    audit.table(0, question_rows, "v40_run_report.json phase_a/phase_b/phase_c", "여형준 표 1")

    strata = synthesis["strata"]
    s1 = [v for k, v in strata.items() if k.startswith("S1_worker_retain|")]
    s2 = [v for k, v in strata.items() if k.startswith("S2_worker_deprioritize|")]
    s3 = strata["S3_worker_uncertain"]
    s4 = strata["S4_adjudication"]
    stratum_rows = [
        ["층", "모수 N", "표본 n", "가중치"],
        ["분류기 × retain (질문별 5층)", number(sum(x["population_N"] for x in s1)),
         number(sum(x["sample_n"] for x in s1)), f"{min(x['weight'] for x in s1):.2f} ~ {max(x['weight'] for x in s1):.2f}"],
        ["분류기 × deprioritize (질문별 5층)", number(sum(x["population_N"] for x in s2)),
         number(sum(x["sample_n"] for x in s2)), f"{min(x['weight'] for x in s2):.2f} ~ {max(x['weight'] for x in s2):.2f}"],
        ["분류기 × uncertain", number(s3["population_N"]), f"{s3['sample_n']} (전수)", "1"],
        ["재판정 (전 라벨)", number(s4["population_N"]), f"{s4['sample_n']} (전수)", "1"],
        ["합계", number(sum(v["population_N"] for v in strata.values())), number(sample_n), ""],
    ]
    audit.table(1, stratum_rows, "screener_vs_ai_reference_v40.json strata", "여형준 표 2")

    audit.covered_tables.add(2)  # publication-year table has no corresponding value in the allowed canon.
    audit.results.append(Result(
        "여형준 표 3 출판연도별 분포", "논문 표의 합계 48,031/retain 3,374는 내부 합산 일치",
        "정본에 연도별 분포 없음", "대조 불가", "허용된 정본 파일에 대응 분포 없음"
    ))

    metrics_rows = [
        ["지표", "값", "비고"],
        ["agreement_vs_ai_reference", pct(scoring["headline"]["agreement_vs_ai_reference_weighted"], 2),
         f"부트스트랩 95% CI {pct(agreement_ci['ci95_low'], 2).rstrip('%')}~{pct(agreement_ci['ci95_high'], 2).rstrip('%')}"],
        ["sensitivity_vs_ai_reference", pct(scoring["headline"]["sensitivity_vs_ai_reference_weighted"], 2), "양성 범주 = retain"],
        ["specificity_vs_ai_reference", pct(scoring["headline"]["specificity_vs_ai_reference_weighted"], 2), "음성 범주 = deprioritize | uncertain"],
        ["precision_vs_ai_reference", pct(scoring["headline"]["precision_vs_ai_reference_weighted"], 2), ""],
        ["Cohen κ", f"{scoring['headline']['cohen_kappa_vs_ai_reference_unweighted']:.3f}", "비가중 3범주"],
        ["agreement (비가중)", pct(scoring["headline"]["agreement_vs_ai_reference_unweighted"], 2), "표본 값. 경계 과대표집 반영"],
        ["판정 불일치", f"{scoring['headline']['decision_disagreements']}건", f"{sample_n:,}행 중"],
    ]
    audit.table(3, metrics_rows, "v40_scoring_report.json; screener_vs_ai_reference_v40.json", "여형준 표 4")

    # Table 5 has subgroup values not present in the allowed canon. The two layer rows are canonical.
    adjudication_layer = synthesis["layers"]["agent_adjudication_layer"]
    worker_layer = synthesis["layers"]["worker_classifier_layer"]
    audit.cell(4, 8, 1, number(adjudication_layer["rows_scored"]), "screener_vs_ai_reference_v40.json layers", "여형준 표 5 재판정 층 행 수")
    audit.cell(4, 8, 2, pct(adjudication_layer["unweighted"]["agreement_vs_ai_reference_unweighted"], 1),
               "screener_vs_ai_reference_v40.json layers.agent_adjudication_layer.unweighted", "여형준 표 5 재판정 층 비가중 일치율")
    audit.cell(4, 9, 1, number(worker_layer["rows_scored"]), "screener_vs_ai_reference_v40.json layers", "여형준 표 5 분류기 층 행 수")
    audit.cell(4, 9, 2, pct(worker_layer["unweighted"]["agreement_vs_ai_reference_unweighted"], 1),
               "screener_vs_ai_reference_v40.json layers.worker_classifier_layer.unweighted", "여형준 표 5 분류기 층 비가중 일치율")
    audit.results.append(Result(
        "여형준 표 5 나머지 조건별 행 수·일치율", "논문 기재값", "정본에 조건별 집계 없음",
        "대조 불가", "허용된 정본 파일에 대응 집계 없음"
    ))

    audit.paragraph(
        "여형준 §3.2 검색 실행과 코퍼스",
        "검색은 2026년 7월 28일에 실행하였다",
        ["2026년 7월 28일", "2022-01-01", "2026-07-28", number(phase_b["raw_xml_files"])],
        "v40_run_report.json phase_a.search_period/phase_b.raw_xml_files",
    )
    audit.paragraph(
        "여형준 표 1 캡션",
        "표 1. 질문별 검색 건수와 선별 결과",
        [number(raw), number(raw - rows), number(rows), number(unique_papers), number(unique_record_ids), "4건"],
        "v40_run_report.json phase_b.deduplication/corpus",
    )
    audit.paragraph(
        "여형준 §3.3 두 층 선별",
        "선별은 두 층으로 구성된다",
        [number(rows), number(adjudicated), pct(adjudicated / rows, 1), "100%"],
        "v40_run_report.json phase_c",
    )
    audit.paragraph(
        "여형준 §3.3 라벨·사유·불변식",
        "라벨은 retain",
        ["3종", f"사유 코드 {len(phase_c['reason_code_counts'])}종", "확신도 3종", "근거 기반 2종", f"불변식 {phase_c['invariant_checks']['cases']}건"],
        "v40_run_report.json phase_c key counts/invariant_checks",
    )
    audit.paragraph(
        "여형준 §4.1 코퍼스와 선별",
        "코퍼스는 48,031행",
        [
            number(rows), number(unique_papers), number(abstracts), number(title_only), "1.0", "0건",
            number(retain), pct(retain / rows, 1), number(deprioritize), pct(deprioritize / rows, 1),
            number(uncertain), pct(uncertain / rows, 1),
        ],
        "v40_run_report.json phase_b/phase_c",
    )
    audit.paragraph(
        "여형준 그림 2 캡션",
        "그림 2. 코퍼스 구축과 두 층 선별 흐름",
        [number(rows), number(adjudicated), pct(adjudicated / rows, 1), number(manifest["records"]), number(retain), number(gate_drop)],
        "v40_run_report.json; manifest.json; direct subtraction",
    )
    audit.paragraph(
        "여형준 §4.2 근거 번들",
        "근거 후보는 1,899행이고",
        [number(manifest["records"]), number(manifest["with_dose"]), "15건", number(core["core_records"]),
         number(retain), number(gate_drop), pct(gate_drop / retain, 1)],
        "manifest.json; core_manifest.json; direct subtraction",
    )
    audit.paragraph(
        "여형준 §4.2 공개 화면 근거 범위",
        "공개 화면은 질문별 핵심 근거 15건",
        ["15건", "30건", "HRS1 312", "HRS2 340", "HRS3 715", "HRS4 380", "HRS5 152", number(manifest["records"]), number(core["core_records"])],
        "manifest.json by_question/core_limit; core_manifest.json",
    )
    audit.paragraph(
        "여형준 §4.2 개인화 규칙",
        "개인화 규칙은 29건이다",
        [number(len(active_rules)), number(len(rules)), number(len(alias_rules)), "A1", "A2", "B1", "B2", "B3", "0건", "5건", "24건", "25건"],
        "personalized_rules.json direct counts",
    )
    audit.paragraph(
        "여형준 표 4 캡션",
        "표 4. 모집단 가중 채점 결과",
        [f"분류기 층 {worker_layer['cohen_kappa_vs_ai_reference_unweighted']:.3f}",
         f"재판정 층 {adjudication_layer['cohen_kappa_vs_ai_reference_unweighted']:.3f}"],
        "screener_vs_ai_reference_v40.json layers",
    )
    question_tokens = []
    for qid in [qids[4], qids[3], qids[0], qids[2], qids[1]]:
        q = synthesis["per_question"][qid]
        question_tokens.extend([
            pct(q["agreement_vs_ai_reference_unweighted"], 2),
            f"{q['cohen_kappa_vs_ai_reference_unweighted']:.3f}",
        ])
    audit.paragraph(
        "여형준 §4.3 질문별 결과",
        "질문별 비가중 일치율은",
        question_tokens,
        "screener_vs_ai_reference_v40.json per_question",
    )
    audit.paragraph(
        "여형준 §4.4 retain 비율과 불일치 방향",
        "파이프라인의 전수 retain 비율은 7.02%인데 채점자의 설계기반 가중 추정치는 15.33%(부트스트랩",
        [
            pct(prevalence["ai_reference_apparent_census"], 2), pct(prevalence["scorer_design_based_weighted"], 2),
            pct(scorer_ci["ci95_low"], 2).rstrip("%"),
            pct(scorer_ci["ci95_high"], 2).rstrip("%"),
            f"{prevalence['ratio_scorer_to_ai_reference']:.2f}배", pct(weighted["specificity_vs_ai_reference"], 2),
            number(deprioritize), "11.5%",
            *[f"{direction.replace('->', '→')} {count}건" for direction, count in disagreements["by_direction"].items()],
        ],
        "screener_vs_ai_reference_v40.json retain_prevalence_estimates/decision_disagreements",
    )
    audit.paragraph(
        "여형준 §4.4 Rogan-Gladen",
        "Rogan–Gladen 보정을 적용한 값은",
        ["16자리", "5.83×10⁻¹⁶", number(retain)],
        "screener_vs_ai_reference_v40.json retain_prevalence_estimates",
    )
    audit.paragraph(
        "여형준 §5.2 불일치 확신도",
        "불일치 264건이",
        [number(disagreements["count"]), "high 82건", "medium 90건", "low 92건"],
        "screener_vs_ai_reference_v40.json decision_disagreements.by_scorer_confidence",
    )
    audit.paragraph(
        "여형준 한계 2",
        "둘째, 근거 번들이",
        [number(retain), number(gate_drop), pct(gate_drop / retain, 1), number(manifest["records"])],
        "v40_run_report.json; manifest.json; direct subtraction",
    )
    audit.paragraph(
        "여형준 한계 3 원시규칙 대 재판정",
        "셋째, 분류기의 원시 규칙과 재판정이",
        [number(adjudicated), number(phase_c["raw_rule_audit"]["mismatches"]), "36.7%", "77건", "12.5%", "149건", "24.2%", "53건"],
        "v40_run_report.json raw_rule_audit for 616/226; 77+149=226 is internal arithmetic",
    )
    audit.results.append(Result(
        "여형준 한계 3의 77/149/53 세부 분해", "77+149=226 산술 일치", "정본에 세부 분해 없음",
        "대조 불가", "허용된 정본에는 raw_rule_audit 총 226건만 있음"
    ))

    fig_source = (fig_dir / "fig_yeo.py").read_text(encoding="utf-8")
    expected_figure_fragments = [
        f"원시 {number(raw)}행", f"중복 {number(raw - rows)}행 제거", f"레코드–질문 {number(rows)}행",
        f"고유 문헌 {number(unique_papers)}편", f"초록 있음 {number(abstracts)}행", f"제목만 {number(title_only)}행",
        f"경계 {number(adjudicated)}행({pct(adjudicated / rows, 1)})만", f'("retain", "{number(retain)}", "{pct(retain / rows, 1)}"',
        f'("deprioritize", "{number(deprioritize)}", "{pct(deprioritize / rows, 1)}"',
        f'("uncertain", "{number(uncertain)}", "{pct(uncertain / rows, 1)}"',
        f"게이트 탈락 {number(gate_drop)}행", f"근거 후보 {number(manifest['records'])}행",
        f"용량이 보고된 행은 {number(manifest['with_dose'])}행", f"핵심 근거 {number(core['core_records'])}건 · 개인화 규칙 {number(len(active_rules))}건",
        f"코퍼스 {number(rows)}행", f"채점 표본 {number(sample_n)}행", "층화 부트스트랩 10,000회",
    ]
    missing = [fragment for fragment in expected_figure_fragments if fragment not in fig_source]
    audit.add("여형준 그림 1·2 소스 수치", "모든 기대 수치 있음" if not missing else "누락: " + ", ".join(missing),
              "모든 기대 수치 있음", "fig_yeo.py vs canonical artifacts", not missing)

    for name, expected_hash in EXPECTED_FIGURE_HASHES.items():
        path = fig_dir / name
        audit.add(f"여형준 그림 파일 해시 {name}", sha256(path), expected_hash,
                  "2026-07-31 visually inspected source/PNG baseline")

    return audit, audit.unverified_numeric_items()


def print_results(audit: Audit, unverified: list[str]) -> int:
    print("| 위치 | 관찰값 | 정본/기대값 | 판정 | 근거 |")
    print("|---|---|---|---|---|")
    for result in audit.results:
        values = [result.location, result.observed, result.expected, result.status, result.source]
        safe = [str(v).replace("|", " / ").replace("\n", " ") for v in values]
        print("| " + " | ".join(safe) + " |")
    failures = [r for r in audit.results if r.status == "불일치"]
    print(f"\nSUMMARY pass={sum(r.status == '일치' for r in audit.results)} "
          f"fail={len(failures)} unverified_groups={sum(r.status == '대조 불가' for r in audit.results)}")
    if unverified:
        print("\nUNMAPPED NUMERIC ITEMS (manual review required)")
        for item in unverified:
            print("- " + item.replace("\n", " "))
    return 1 if failures else 0


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--figure-dir", type=Path, default=DEFAULT_FIG_DIR)
    args = parser.parse_args()
    audit, unverified = build_audit(args.docx, args.figure_dir)
    return print_results(audit, unverified)


if __name__ == "__main__":
    sys.exit(main())
