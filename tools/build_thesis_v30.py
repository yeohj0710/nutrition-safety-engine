#!/usr/bin/env python3
"""검증된 manifest에서 protocol v3.0 한국어 학위논문 DOCX/Markdown을 생성한다.

모든 수치는 manifest에서 읽는다. 본문에 숫자를 직접 적지 않는다.
편집 기준은 research/thesis/etc/template_v21/reference_v21.docx 의
쪽 구성·여백·제목 체계·Pretendard 글꼴을 그대로 따른다.
"""
from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "research/thesis"
OUT.mkdir(exist_ok=True)
DOCX = OUT / "thesis_v30.docx"
MD = OUT / "thesis_v30_ko.md"
AUTHOR = "여형준"
SUBMISSION_DATE = "2026년 7월"

P = {
    "picos": ROOT / "research/searches_v3/ai_picos/picos_definition.json",
    "picos_prompt": ROOT / "research/searches_v3/ai_picos/prompt.txt",
    "corpus_manifest": ROOT / "data/curated_v3/corpus_manifest.json",
    "corpus": ROOT / "data/curated_v3/evidence_map.csv",
    "screen_manifest": ROOT / "research/screening/v30_agent/manifest.json",
    "screen_prompt": ROOT / "research/screening/v30_agent/prompts/screening_prompt.md",
    "screen_output": ROOT / "data/curated_v3/llm_screening_classifications.csv",
    "reference_prompt": ROOT
    / "research/validation/screening_ai_reference_v3/prompts/reference_picos_prompt.md",
    "reference": ROOT / "research/synthesis/screener_vs_ai_reference_v3.json",
    "review_manifest": ROOT / "research/systematic_review_v30/manifest.json",
    "core_manifest": ROOT / "research/systematic_review_v30/core_manifest.json",
    "track_comparison": ROOT / "research/synthesis/picos_track_comparison.json",
}

PICOS = json.loads(P["picos"].read_text(encoding="utf-8"))
CM = json.loads(P["corpus_manifest"].read_text(encoding="utf-8"))
SM = json.loads(P["screen_manifest"].read_text(encoding="utf-8"))
REF = json.loads(P["reference"].read_text(encoding="utf-8"))
RM = json.loads(P["review_manifest"].read_text(encoding="utf-8"))
CORE = json.loads(P["core_manifest"].read_text(encoding="utf-8"))
TC = json.loads(P["track_comparison"].read_text(encoding="utf-8"))

QUESTIONS = {q["question_id"]: q for q in PICOS["questions"]}
QORDER = [r["question_id"] for r in CM["search"]["question_runs"]]
RUNS = {r["question_id"]: r for r in CM["search"]["question_runs"]}
WM = REF["weighted_metrics"]
BOOT = REF["bootstrap"]
CORPUS = REF["corpus"]
ROUNDS = REF["rounds"]
GATE = RM["llm_gate"]
SOURCE_LABEL = {"pubmed_only": "PubMed 단일 자료원"}.get(
    CM["source_constraint"], CM["source_constraint"]
)


def sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def pct(value: float, digits: int = 1) -> str:
    return f"{value * 100:.{digits}f}%"


def num(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _no_split(row) -> None:
    """행이 쪽 경계에서 잘려 빈 칸만 남는 것을 막는다."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    _no_split(t.rows[0])
    for row in rows:
        new_row = t.add_row()
        for i, value in enumerate(row):
            new_row.cells[i].text = str(value)
        _no_split(new_row)
    return t


# ---------------------------------------------------------------- 본문 구성
picos_rows = [
    [qid, QUESTIONS[qid]["P"], QUESTIONS[qid]["I"], QUESTIONS[qid]["O"]]
    for qid in QORDER
]

search_rows = [
    [
        qid,
        QUESTIONS[qid]["P"],
        f"{RUNS[qid]['hit_count']:,}",
        f"{RUNS[qid]['retrieved_count']:,}",
    ]
    for qid in QORDER
]

screen_rows = [
    [
        qid,
        f"{SM['distribution_by_question'][qid]['retain']:,}",
        f"{SM['distribution_by_question'][qid]['deprioritize']:,}",
        f"{SM['distribution_by_question'][qid]['uncertain']:,}",
    ]
    for qid in QORDER
]

gate_rows = [[qid, f"{RM['by_question'][qid]:,}"] for qid in QORDER]

stratum_rows = [
    [
        s["stratum_id"],
        f"{s['frame_size']:,}",
        f"{s['sample_size']:,}",
        num(s["weight"], 3),
    ]
    for s in REF["sample"]["strata"]
]

metric_rows = [
    [
        "sensitivity_vs_ai_reference",
        num(WM["sensitivity_vs_ai_reference"]),
        f"{num(BOOT['sensitivity_vs_ai_reference_ci95'][0])} – {num(BOOT['sensitivity_vs_ai_reference_ci95'][1])}",
    ],
    [
        "specificity_vs_ai_reference",
        num(WM["specificity_vs_ai_reference"]),
        f"{num(BOOT['specificity_vs_ai_reference_ci95'][0])} – {num(BOOT['specificity_vs_ai_reference_ci95'][1])}",
    ],
    ["agreement_vs_ai_reference", num(WM["agreement_vs_ai_reference"]), "—"],
    [
        "보정 전 겉보기 retain 비율",
        num(CORPUS["apparent_retain_share"]),
        "—",
    ],
    [
        "Rogan–Gladen 보정 retain 비율",
        num(CORPUS["rogan_gladen_corrected_retain_share"]),
        f"{num(BOOT['corrected_retain_share_ci95'][0])} – {num(BOOT['corrected_retain_share_ci95'][1])}",
    ],
]

round_rows = [
    [
        key.replace("round", "라운드 ").replace("_vs_ ", " 대 "),
        num(value["cohen_kappa"], 4),
        pct(value["raw_agreement"], 1),
        f"{value['axis_cells_differing']:,}/{value['axis_cells_compared']:,}",
    ]
    for key, value in ROUNDS["pairwise"].items()
]

abstract = (
    f"본 연구는 고위험 임상상황에서의 영양보충제 안전성 문헌을 대상으로, 연구 질문(PICOS)과 "
    f"PubMed 검색식을 사람이 아니라 인공지능이 정의하고, 문헌 선별과 참조 판정까지 인공지능 "
    f"에이전트가 직접 수행하는 완전 자율 트랙(protocol v3.0)을 구축하고 그 산출물을 보고한다. "
    f"질문은 {CM['picos']['question_count']}개이며, 자료원은 {SOURCE_LABEL}으로 한정하였다. "
    f"검색으로 확보한 코퍼스는 레코드-질문 단위 {CM['corpus']['row_count']:,}건(고유 문헌 "
    f"{CM['corpus']['unique_record_count']:,}건)이고, 이 가운데 초록이 있는 행은 "
    f"{CM['corpus']['observability_distribution']['abstract_available']:,}건, 제목만 있는 행은 "
    f"{CM['corpus']['observability_distribution']['title_only']:,}건이었다. "
    f"에이전트는 이 {SM['row_count']:,}행을 빠짐없이 직접 판정하여 커버리지 {pct(SM['coverage'], 0)}를 "
    f"달성했고, 판정 분포는 retain {SM['distribution']['retain']:,}건, deprioritize "
    f"{SM['distribution']['deprioritize']:,}건, uncertain {SM['distribution']['uncertain']:,}건이었다. "
    f"선별 품질을 가늠하기 위해 층화 무작위 표본 {REF['sample']['sample_size']}건을 뽑아 "
    f"P·I·C·O·S 축을 라운드 {ROUNDS['count']}회로 독립 채점하고 코드 규칙으로 참조 라벨을 도출했다. "
    f"층화 가중 결과 sensitivity_vs_ai_reference는 {num(WM['sensitivity_vs_ai_reference'])}, "
    f"specificity_vs_ai_reference는 {num(WM['specificity_vs_ai_reference'])}, "
    f"agreement_vs_ai_reference는 {num(WM['agreement_vs_ai_reference'])}였다. "
    f"Rogan–Gladen 보정을 적용하면 코퍼스 수준 retain 규모는 겉보기 "
    f"{CORPUS['apparent_retain_count']:,}건에서 "
    f"{round(CORPUS['rogan_gladen_corrected_retain_count']):,}건(95% CI "
    f"{round(BOOT['corrected_retain_count_ci95'][0]):,}–"
    f"{round(BOOT['corrected_retain_count_ci95'][1]):,}건)으로 낮아진다. "
    f"이 참조 판정은 사람이 만든 gold standard가 아니라 같은 에이전트가 블라인드 상태에서 다른 "
    f"과업 정의로 수행한 내부 참조(ai_reference_standard)이므로, 위 값은 임상적 정확도가 아니라 "
    f"교차 확인(ai_cross_checked) 결과로만 해석해야 한다. 본 연구의 모든 단계에서 사람의 연구 "
    f"의사결정은 {CM['human_decisions']}건이다."
)

sections = [
    (
        "1. 서론",
        [
            (
                "1.1 연구 배경",
                [
                    "영양보충제는 처방 없이 구할 수 있지만, 수술 전후·만성콩팥병·임신·간질환·항응고 치료처럼 "
                    "위해가 커질 수 있는 상황에서는 성분의 평균적인 효과보다 노출량, 병용약, 기저질환, 관찰된 "
                    "이상반응을 함께 확인해야 한다. 이런 확인에 쓰이는 근거는 무작위 대조시험부터 증례보고까지 "
                    "여러 형태로 흩어져 있어, 무엇을 먼저 읽어야 하는지 정하는 일 자체가 하나의 과업이 된다.",
                    "문헌을 선별하는 작업은 전통적으로 두 명 이상의 사람이 독립적으로 수행하고 이견을 조정한다. "
                    "그러나 인력과 원문 접근을 확보하지 못한 연구 환경에서는 이 절차가 성립하지 않는다. "
                    "이때 선택지는 연구를 멈추거나, 자동화가 어디까지 할 수 있고 어디부터 할 수 없는지를 "
                    "명시적으로 드러내는 설계를 만드는 것이다.",
                ],
            ),
            (
                "1.2 문제 제기",
                [
                    "선행 트랙(protocol v2.1)은 사람이 정의한 질문과 검색식을 그대로 두고 자동 분류만 비교했다. "
                    "이 설계에서는 인공지능이 실제로 연구 질문을 세우고 검색 전략을 짜는 단계까지 담당했을 때 "
                    "무엇이 달라지는지 관찰할 수 없다. 또한 자동 분류의 품질을 가늠할 기준이 없어, 분류 결과를 "
                    "그대로 근거 규모로 읽으면 과대추정이 발생한다.",
                    "본 연구는 그 두 가지 공백을 겨냥한다. 첫째, 질문 정의와 검색식 작성부터 인공지능이 수행하는 "
                    "독립 트랙을 만든다. 둘째, 같은 코퍼스에 대해 별도의 과업 정의로 참조 판정을 생성하고, "
                    "그 참조와의 대조 결과로 분류기의 겉보기 규모를 보정한다.",
                ],
            ),
            (
                "1.3 연구 목적과 질문",
                [
                    "연구 목적은 네 가지다. 첫째, 인공지능이 정의한 고위험 임상상황 질문과 PubMed 검색식으로 "
                    "선행 트랙과 분리된 코퍼스를 구축한다. 둘째, 그 코퍼스 전량을 에이전트가 직접 선별한다. "
                    "셋째, 층화 표본에 대해 축 단위 참조 판정을 만들고 층화 가중과 Rogan–Gladen 보정으로 "
                    "코퍼스 수준 retain 규모를 추정한다. 넷째, 그 결과를 공개 화면의 근거로 연결하고 "
                    "재현 경로를 해시로 고정한다.",
                    "질문은 표 1과 같이 다섯 개이며, 분석의 기본 단위는 개별 연구가 아니라 특정 문헌이 특정 "
                    "질문에서 검색된 레코드-질문 단위다.",
                ],
            ),
            (
                "1.4 연구 범위",
                [
                    "본 논문은 임상 권고를 제시하지 않는다. 효과크기를 통합하거나 위험비를 계산하지 않고, "
                    "비뚤림 위험이나 근거수준을 판정하지 않으며, PRISMA 흐름도의 최종 포함·제외 건수를 "
                    "보고하지 않는다. 선별 라벨은 retain / deprioritize / uncertain이며 사람의 포함·제외 "
                    "판정이 아니다.",
                ],
            ),
        ],
    ),
    (
        "2. 이론적·방법론적 배경",
        [
            (
                "2.1 완전 자율 트랙의 정의",
                [
                    "여기서 완전 자율이란 질문 정의, 검색식 작성, 검색 실행, 선별, 참조 판정, 번역, 집필까지 "
                    "인공지능이 수행하고 사람의 연구 의사결정이 개입하지 않는 상태를 말한다. 자율성은 정확도의 "
                    "주장이 아니라 절차의 기술이며, 그래서 이 논문은 자율 트랙의 산출물을 정확도가 아니라 "
                    "재현 가능성과 내부 일관성으로 평가한다.",
                ],
            ),
            (
                "2.2 AI 참조표준의 위치",
                [
                    "사람 gold standard가 없는 상황에서 분류기 성능을 평가하려면 비교 대상이 필요하다. 본 연구는 "
                    "같은 에이전트가 블라인드 상태에서, 선별과 다른 과업 정의(축 단위 채점)로 만든 판정을 "
                    "참조로 삼았다. 이를 ai_reference_standard로 부르고, 대조 결과는 진실 정확도가 아니라 "
                    "ai_cross_checked 결과로 표기한다.",
                    "이 구분은 용어에서도 지킨다. 민감도·특이도·정확도를 단독으로 쓰지 않고 "
                    "sensitivity_vs_ai_reference, specificity_vs_ai_reference, agreement_vs_ai_reference로 "
                    "적는다.",
                ],
            ),
            (
                "2.3 유병률 보정의 필요",
                [
                    "불완전한 검사로 관찰한 양성 비율은 참값이 아니다. 검사의 민감도와 특이도를 알면 "
                    "Rogan–Gladen 보정으로 참 비율을 추정할 수 있다. 본 연구는 분류기를 검사로, 참조 판정을 "
                    "기준으로 두고 이 보정을 적용해, 선별 결과를 그대로 근거 규모로 읽을 때 생기는 과대추정을 "
                    "수치로 드러냈다.",
                ],
            ),
        ],
    ),
    (
        "3. 연구 방법",
        [
            (
                "3.1 질문과 검색식의 AI 정의",
                [
                    f"연구 질문과 PubMed 검색식은 AI 에이전트가 정의하였고 이 단계에서 사람의 "
                    f"판정은 {PICOS['generated_by']['human_decisions']}건이다. 정의를 수행한 실행 주체는 "
                    f"{PICOS['generated_by']['model']}이며, 뒤이은 선별·참조 판정·번역·집필을 수행한 "
                    f"에이전트와는 다른 실행 주체다. 정의 프롬프트는 "
                    f"{P['picos_prompt'].relative_to(ROOT).as_posix()}에 두고 "
                    f"SHA-256 {CM['picos']['prompt_sha256']}로 고정했다. "
                    f"정의 결과 파일의 SHA-256은 {CM['picos']['sha256']}다.",
                    f"질문 선정의 근거는 다음과 같이 기록되어 있다. “{PICOS['derivation']}” "
                    f"정의 단계의 입력에서는 {', '.join(PICOS['input']['excluded_inputs'])}를 명시적으로 "
                    f"배제하여, 선행 트랙의 질문·검색식·결과 수치가 새 질문 정의에 흘러들어가지 않도록 했다.",
                ],
            ),
            (
                "3.2 검색 실행과 코퍼스",
                [
                    f"검색은 {RUNS[QORDER[0]]['database'].upper()}를 대상으로 "
                    f"{RUNS[QORDER[0]]['executed_at']}에 실행하였다. 질문별 hit 수와 확보 건수는 표 2와 같고, "
                    f"합계 hit 수는 {CM['search']['combined_hit_count']:,}건이다.",
                    f"표준화한 코퍼스는 {CM['corpus']['path']}이며 레코드-질문 단위 "
                    f"{CM['corpus']['row_count']:,}건, 고유 문헌 {CM['corpus']['unique_record_count']:,}건, "
                    f"SHA-256 {CM['corpus']['sha256']}다. 관찰 가능성은 초록 보유 "
                    f"{CM['corpus']['observability_distribution']['abstract_available']:,}건, 제목만 "
                    f"{CM['corpus']['observability_distribution']['title_only']:,}건으로 나뉜다.",
                ],
            ),
            (
                "3.3 에이전트 직접 선별",
                [
                    f"선별 프롬프트는 {SM['prompt_path']}에 동결하고 SHA-256 {SM['prompt_sha256']}로 "
                    f"기록했다. 선별은 에이전트가 배치를 직접 읽고 판정하는 방식이며, 선별을 위한 "
                    f"별도 모델 호출은 {SM['model_invocations']}회, 외부 API 호출은 "
                    f"{SM['external_api_calls']}회, 사람의 판정은 {SM['human_decisions']}건이다. "
                    f"즉 에이전트가 배치를 읽고 모든 행을 직접 판정했다.",
                    f"판정은 append-only 체크포인트({SM['checkpoint_path']}, SHA-256 "
                    f"{SM['checkpoint_sha256']})에 누적하고, 요청한 레코드-질문 쌍이 정확히 한 번씩 "
                    f"돌아왔는지 반복 검증하여 커버리지 {pct(SM['coverage'], 0)}에 도달한 뒤 확정했다. "
                    f"최종 산출물은 {SM['output_path']}(SHA-256 {SM['output_sha256']})다.",
                    f"초록이 없는 {SM['by_evidence_basis']['title_only']['rows']:,}행은 근거 형태를 "
                    f"title_only로 고정하고 confidence 상한을 low로 두었으며 사유 코드에 "
                    f"insufficient_abstract를 강제했다. 판정 실패나 파싱 오류라는 상태는 두지 않았고, "
                    f"판단이 어려운 행은 uncertain으로 남겼다.",
                ],
            ),
            (
                "3.4 AI 참조표준의 생성과 채점",
                [
                    f"선별 결과를 층으로 삼아 층화 무작위 표본 {REF['sample']['sample_size']}건을 뽑았다. "
                    f"표본 시드는 {REF['sample']['sample_seed']}이며 층별 프레임 크기·표본 수·가중치는 "
                    f"표 5와 같다. 블라인드 파일에는 선별 라벨·confidence·사유 코드·배치 번호를 넣지 않았다.",
                    f"참조 채점 프롬프트는 선별 프롬프트와 다른 과업을 정의한다. 주제 적합성을 통째로 묻지 않고 "
                    f"P·I·C·O·S 다섯 축을 각각 채점한 뒤, 종합 라벨은 사람이 아니라 코드의 명시적 규칙이 "
                    f"도출한다. 프롬프트는 {P['reference_prompt'].relative_to(ROOT).as_posix()}에 동결했고 "
                    f"SHA-256은 {REF['inputs']['prompt_sha256']}다.",
                    f"채점은 라운드마다 행 순서를 독립 무작위화하여 {ROUNDS['count']}회 수행했다(시드 "
                    f"{', '.join(str(v) for v in ROUNDS['seeds'].values())}). 세 라운드의 다수결로 참조 라벨을 "
                    f"정하고, 세 라벨이 모두 다르면 unresolved로 남기기로 했으나 실제 unresolved는 "
                    f"{REF['sample']['unresolved_excluded']}건이었다. 분석에 사용한 단위는 "
                    f"{REF['sample']['analysed_units']}건이다.",
                    f"참조 채점 역시 {REF['execution']['scorer']} 방식이며 모델 호출 "
                    f"{REF['execution']['model_invocations']}회, 외부 API 호출 "
                    f"{REF['execution']['external_api_calls']}회, 사람의 판정 "
                    f"{REF['execution']['human_decisions']}건이다.",
                ],
            ),
            (
                "3.5 통계 처리",
                [
                    "층별 추출률이 다르므로 단순 평균을 쓰지 않고 층화 가중치를 적용했다. 가중치는 층 프레임 "
                    "크기를 층 표본 수로 나눈 값이다.",
                    f"코퍼스 수준 retain 규모는 Rogan–Gladen 보정으로 추정했다. 신뢰구간은 층 안에서 "
                    f"복원 추출하는 층화 부트스트랩 {BOOT['iterations']:,}회의 백분위수 방식으로 계산했고 "
                    f"시드는 {BOOT['seed']}다. 통계 함수(층화 가중, Rogan–Gladen, 부트스트랩 CI, 백분위수, "
                    f"라벨 도출 규칙, 층별 배분, kappa)에는 단위 시험을 붙였다.",
                ],
            ),
            (
                "3.6 근거 번들과 공개 화면 연결",
                [
                    f"선별을 통과한 문헌은 정규식 PICOS 추출 게이트와 결합하여 근거 번들로 만들었다. "
                    f"정규식 게이트 통과 {GATE['regex_passed']:,}행 가운데 선별 라벨이 retain이 아닌 "
                    f"{GATE['dropped_by_llm']:,}행이 제외되어 {GATE['kept']:,}행이 남았다. "
                    f"질문당 상한 {RM['core_limit_per_question']}건을 적용해 핵심 근거 "
                    f"{CORE['core_records']:,}건을 선정했다.",
                    "핵심 근거의 한국어 핵심소견 번역은 외부·로컬 번역 모델을 호출하지 않고 에이전트가 직접 "
                    "작성했다. 번역기는 원문의 숫자·단위·증감 방향이 보존되었는지 자동 검증하며, 검증에 "
                    "실패하면 번들 생성이 중단된다.",
                ],
            ),
            (
                "3.7 선행 트랙과의 비교 절차",
                [
                    "v2.1 트랙은 사람이 질문과 검색식을 정의한 비교 트랙으로 남기고 덮어쓰지 않았다. "
                    "두 트랙의 질문 포괄 관계, 검색식 용어 Jaccard, MeSH 사용, hit 수, PMID 교집합을 "
                    "별도 도구로 산출했다.",
                ],
            ),
        ],
    ),
    (
        "4. 연구 결과",
        [
            (
                "4.1 코퍼스 구성",
                [
                    f"코퍼스는 레코드-질문 단위 {CM['corpus']['row_count']:,}건이며 자료원은 "
                    f"{', '.join(f'{k.upper()} {v:,}건' for k, v in CM['corpus']['source_distribution'].items())}이다. "
                    f"고유 문헌은 {CM['corpus']['unique_record_count']:,}건으로, 한 문헌이 둘 이상의 질문에서 "
                    f"검색된 경우가 있다. 이 값은 중복을 확정 제거한 연구 수가 아니다.",
                ],
            ),
            (
                "4.2 선별 결과",
                [
                    f"전체 {SM['classified']:,}행의 판정 분포는 retain {SM['distribution']['retain']:,}건, "
                    f"deprioritize {SM['distribution']['deprioritize']:,}건, uncertain "
                    f"{SM['distribution']['uncertain']:,}건이었다. 질문별 분포는 표 3과 같다.",
                    f"근거 형태별로 보면 초록이 있는 {SM['by_evidence_basis']['abstract']['rows']:,}행에서는 "
                    f"uncertain이 {SM['by_evidence_basis']['abstract']['distribution']['uncertain']:,}건에 "
                    f"그쳤으나, 제목만 있는 {SM['by_evidence_basis']['title_only']['rows']:,}행에서는 "
                    f"{SM['by_evidence_basis']['title_only']['distribution']['uncertain']:,}건이었다. "
                    f"확신도 분포는 high {SM['confidence_distribution']['high']:,}건, medium "
                    f"{SM['confidence_distribution']['medium']:,}건, low "
                    f"{SM['confidence_distribution']['low']:,}건이다.",
                ],
            ),
            (
                "4.3 참조 판정과의 대조",
                [
                    f"다수결 참조 라벨의 분포와 라운드 간 일치는 표 6에 정리했다. 라운드 전체가 동일한 라벨을 "
                    f"준 비율은 {pct(ROUNDS['unanimous_share'], 1)}였다.",
                    f"층화 가중 결과 sensitivity_vs_ai_reference는 {num(WM['sensitivity_vs_ai_reference'])}, "
                    f"specificity_vs_ai_reference는 {num(WM['specificity_vs_ai_reference'])}, "
                    f"agreement_vs_ai_reference는 {num(WM['agreement_vs_ai_reference'])}였다(표 7). "
                    f"즉 참조가 retain으로 본 문헌은 분류기가 거의 놓치지 않았으나, 참조가 deprioritize로 본 "
                    f"문헌의 상당수를 분류기는 retain으로 남겼다.",
                    f"그 결과 겉보기 retain 비율 {num(CORPUS['apparent_retain_share'])}"
                    f"({CORPUS['apparent_retain_count']:,}건)는 Rogan–Gladen 보정 후 "
                    f"{num(CORPUS['rogan_gladen_corrected_retain_share'])}"
                    f"({round(CORPUS['rogan_gladen_corrected_retain_count']):,}건, 95% CI "
                    f"{round(BOOT['corrected_retain_count_ci95'][0]):,}–"
                    f"{round(BOOT['corrected_retain_count_ci95'][1]):,}건)으로 낮아진다. "
                    f"선별 결과를 그대로 근거 규모로 읽으면 과대추정이 발생한다는 뜻이다.",
                    f"불일치 사례는 분류기 retain·참조 deprioritize "
                    f"{REF['disagreement_examples']['counts']['classifier_positive_reference_negative']:,}건, "
                    f"분류기 deprioritize·참조 retain "
                    f"{REF['disagreement_examples']['counts']['classifier_negative_reference_positive']:,}건이며 "
                    f"실제 제목과 함께 산출물에 보존했다.",
                ],
            ),
            (
                "4.4 근거 번들",
                [
                    f"정규식 게이트와 선별 라벨을 함께 적용해 {GATE['kept']:,}행이 남았고 질문별 분포는 표 4와 "
                    f"같다. 핵심 근거는 {CORE['core_records']:,}건이며 한국어 번역도 같은 수만큼 작성했다. "
                    f"근거 번들에는 효과 수치를 임의로 채우지 않아 효과 상태는 전부 "
                    f"{list(RM['effect_status'])[0]}로 남겼다.",
                ],
            ),
            (
                "4.5 선행 트랙과의 비교",
                [
                    f"질문 포괄성 검토 결과 v2.1의 "
                    f"{', '.join(TC['question_level_coverage']['covered_legacy_question_ids'])}는 v3.0 질문이 "
                    f"포괄하고, {', '.join(TC['question_level_coverage']['uncovered_legacy_question_ids'])}는 "
                    f"포괄하지 못했다. 기록된 판단은 다음과 같다. "
                    f"“{TC['question_level_coverage']['finding']}”",
                    f"검색 결과 수준에서 두 트랙의 PMID 교집합은 "
                    f"{TC['result_level']['track_union_pmid_comparison']['intersection_count']:,}건, "
                    f"v2.1 전용은 "
                    f"{TC['result_level']['track_union_pmid_comparison']['v2_only_count']:,}건, "
                    f"v3.0 전용은 "
                    f"{TC['result_level']['track_union_pmid_comparison']['v3_only_count']:,}건이다. "
                    f"두 트랙은 사실상 다른 문헌 집합을 보고 있다.",
                ],
            ),
        ],
    ),
    (
        "5. 고찰",
        [
            (
                "5.1 주요 결과의 의미",
                [
                    "가장 중요한 결과는 선별 통과 건수 자체가 아니라, 그 건수를 그대로 근거 규모로 읽으면 안 "
                    "된다는 점을 같은 트랙 안에서 수치로 보인 것이다. 민감도는 높고 특이도는 중간 수준이라는 "
                    "패턴은, 자동 선별이 놓치는 문헌을 줄이는 쪽으로 치우쳐 있음을 뜻한다. 탐색 단계에서는 "
                    "바람직한 성질이지만, 보정 없이 규모를 보고하면 근거가 실제보다 두껍게 보인다.",
                    "라운드 간 일치가 높게 나온 것은 축 정의가 안정적으로 적용됐다는 뜻이지만, 채점자가 한 "
                    "명이므로 이는 서로 다른 채점자 사이의 일치가 아니라 동일 채점자의 재검사 안정성이다. "
                    "실제로 두 번째와 세 번째 라운드는 축 셀이 하나도 다르지 않았고, 이는 독립적인 3차 판정이 "
                    "아니라 2차 판정의 재현으로 보아야 한다. 정보량이 있는 값은 기준이 굳어지는 중이던 첫 "
                    "라운드와의 비교뿐이다.",
                ],
            ),
            (
                "5.2 방법론적 기여",
                [
                    "질문 정의부터 집필까지 사람의 연구 의사결정 없이 수행한 트랙을, 각 단계의 프롬프트 해시와 "
                    "입력·출력 해시로 고정해 재현 가능하게 남겼다. 자율성을 성능 주장으로 쓰지 않고 절차 기술로 "
                    "다룬 점, 그리고 그 절차가 만든 편향을 같은 트랙 안에서 보정해 보인 점이 기여다.",
                    "참조 판정을 축 단위로 나누고 종합 라벨을 코드 규칙에 맡긴 설계는, 같은 모델이 같은 기준을 "
                    "두 번 적용해 일치도를 부풀리는 문제를 일부 완화한다. 다만 완전히 없애지는 못한다.",
                ],
            ),
            (
                "5.3 실무적 함의",
                [
                    "자동 선별 결과를 근거 규모로 보고할 때는 참조 대조와 유병률 보정을 함께 제시해야 한다. "
                    "보정 없이 제시된 선별 통과 건수는 후속 검토 인력 산정과 근거 충분성 판단을 모두 왜곡한다.",
                ],
            ),
            (
                "5.4 후속 연구",
                [
                    "다음 단계는 사람 두 명 이상이 독립적으로 같은 층화 표본을 판정해 기준 정답을 만들고, 그것과 "
                    "AI 참조표준을 다시 대조하는 것이다. 그때 비로소 민감도·특이도를 임상적 의미로 해석할 수 "
                    "있다. 또한 자료원을 PubMed 밖으로 넓히고, 원문 접근이 가능한 부분집합에서 초록 기반 판정과 "
                    "원문 기반 판정이 얼마나 달라지는지 확인할 필요가 있다.",
                ],
            ),
        ],
    ),
    (
        "6. 연구의 한계",
        [
            (
                "6.1 참조표준의 성격",
                [
                    "ai_reference_standard는 사람이 만든 gold standard가 아니라 같은 에이전트가 다른 과업 "
                    "정의로 수행한 내부 참조 판정이다. 따라서 본문에 보고한 sensitivity_vs_ai_reference와 "
                    "specificity_vs_ai_reference는 진실 정확도가 아니며, 임상적 정확도로 인용해서는 안 된다.",
                ],
            ),
            (
                "6.2 독립성의 부분적 한계",
                [
                    "분류기와 참조 판정을 같은 모델이 수행했다. 블라인드 처리와 과업 정의 분리, 라운드별 순서 "
                    "무작위화로 독립성을 높였으나, 같은 모델이 공유하는 판단 성향까지 제거하지는 못한다. "
                    "그러므로 두 판정의 일치는 서로 다른 관찰자의 합의보다 약한 증거다.",
                    "라운드 2와 3의 축 판정이 완전히 동일했다는 사실은 이 한계를 직접 보여 준다. 반복 채점이 "
                    "독립적인 재판정이 아니라 앞선 판정의 재현이 될 수 있다.",
                ],
            ),
            (
                "6.3 사람의 연구 의사결정 부재",
                [
                    f"질문 정의, 검색식 작성, 선별, 참조 판정, 번역, 집필 전 단계에서 사람의 연구 의사결정은 "
                    f"{CM['human_decisions']}건이다. 이는 설계상의 선택이지만, 동시에 임상적 타당성을 검증할 "
                    f"주체가 없다는 뜻이기도 하다. 본 연구의 결과는 사람 검토를 대체하지 않는다.",
                ],
            ),
            (
                "6.4 단일 자료원",
                [
                    f"자료원은 {SOURCE_LABEL}({CM['source_constraint']})으로 한정했다. Embase, Scopus, Web of Science, 국내 "
                    f"데이터베이스, 임상시험 등록자료는 포함하지 않았다. 따라서 이 코퍼스는 가능한 모든 근거를 "
                    f"담은 목록이 아니다.",
                ],
            ),
            (
                "6.5 원문 미접근",
                [
                    f"판정과 추출은 제목과 초록 범위에서 이루어졌다. 근거 번들의 원문 위치 확보 건수는 "
                    f"{RM['with_fulltext_locator']}건이며, 접근하지 못한 원문의 내용을 근거로 해석을 넓히지 "
                    f"않았다. 초록에 없는 정보는 없는 것으로 두었다.",
                ],
            ),
        ],
    ),
    (
        "7. 결론",
        [
            (
                "",
                [
                    f"본 연구는 인공지능이 질문과 검색식을 정의하고 코퍼스 "
                    f"{CM['corpus']['row_count']:,}행 전량을 직접 선별한 뒤, 층화 표본 "
                    f"{REF['sample']['sample_size']}건에 대한 축 단위 참조 판정으로 그 선별을 교차 확인하는 "
                    f"완전 자율 트랙을 구축했다. 커버리지는 {pct(SM['coverage'], 0)}, 사람의 연구 의사결정은 "
                    f"{CM['human_decisions']}건이다.",
                    f"핵심 결과는 겉보기 retain {CORPUS['apparent_retain_count']:,}건이 보정 후 "
                    f"{round(CORPUS['rogan_gladen_corrected_retain_count']):,}건으로 낮아진다는 것이다. "
                    f"자동 선별의 산출량은 그대로 근거의 양이 아니며, 참조 대조와 보정을 함께 보고할 때에만 "
                    f"해석 가능한 수치가 된다.",
                    "이 논문은 임상 권고도, 사람이 선별한 체계적 문헌고찰의 결론도 아니다. 의의는 자율 트랙이 "
                    "무엇을 할 수 있고 어디에서 사람이 필요한지를 같은 자료 안에서 구분해 보이고, 모든 수치를 "
                    "원자료와 코드에 해시로 연결해 둔 데 있다.",
                ],
            ),
        ],
    ),
]

# ---------------------------------------------------------------- DOCX 생성
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.top_margin = sec.bottom_margin = Inches(0.85)
sec.left_margin = sec.right_margin = Inches(1)
for name, size, color, font in [
    ("Normal", 10.5, "222222", "Pretendard"),
    ("Title", 20, "15314B", "Pretendard ExtraBold"),
    ("Heading 1", 16, "1F4E79", "Pretendard Bold"),
    ("Heading 2", 13, "1F4E79", "Pretendard SemiBold"),
    ("Heading 3", 11.5, "365F7D", "Pretendard Medium"),
]:
    style = doc.styles[name]
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:ascii"), font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.line_spacing = 1.55 if name == "Normal" else 1.2
    style.paragraph_format.space_after = Pt(7)

sec.header.paragraphs[0].text = "고위험 임상상황 영양보충제 안전성 근거지도 (protocol v3.0)"
add_page_number(sec.footer.paragraphs[0])

TITLE_KO = (
    "고위험 임상상황의 영양보충제 안전성 문헌에 대한\n"
    "완전 AI 자율 근거지도 구축과\nAI 참조표준 기반 선별 교차 확인"
)
TITLE_EN = (
    "Fully AI-Autonomous Evidence Mapping of Dietary Supplement Safety in "
    "High-Risk Clinical Contexts with AI-Reference Cross-Checking of Screening"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(90)
p.add_run("학위논문").bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(35)
run = p.add_run(TITLE_KO)
run.bold = True
run.font.size = Pt(20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(TITLE_EN).italic = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
p.add_run(f"{AUTHOR}\n{SUBMISSION_DATE}").bold = True

doc.add_page_break()
doc.add_heading("제출 정보", 1)
doc.add_paragraph(
    "이 문서는 여형준의 졸업논문 제출용 편집 기준본이다. 연구 내용과 검증된 수치는 본문에 "
    "반영되어 있으며, 아래 행정 정보는 소속 기관의 공식 양식에 맞춰 제출 전에 입력한다."
)
table(
    doc,
    ["항목", "내용"],
    [
        ["연구자", AUTHOR],
        ["소속 대학", "[입력]"],
        ["학과·전공", "[입력]"],
        ["학번", "[입력]"],
        ["지도교수", "[입력]"],
        ["학위명", "[입력]"],
        ["최종 제출일", "[입력]"],
    ],
)
doc.add_paragraph(
    "연구 성격: 인공지능이 연구 질문과 검색식을 정의하고 선별·참조 판정을 직접 수행한 완전 자율 "
    "근거지도 연구. 사람의 독립 이중선별, 최종 RoB·GRADE, 임상 권고를 완료한 체계적 문헌고찰로 "
    "표시하지 않는다."
)

doc.add_page_break()
doc.add_heading("국문초록", 1)
doc.add_paragraph(abstract)
doc.add_paragraph(
    "주요어: 영양보충제, 고위험 임상상황, 근거지도, 인공지능 자율 연구, AI 참조표준, "
    "Rogan–Gladen 보정, 출처 계보"
)

doc.add_heading("Abstract", 1)
doc.add_paragraph(
    "This study built a fully AI-autonomous evidence-mapping track in which the research "
    "questions and PubMed search strategies were defined by AI, and every record-question row "
    f"of the resulting corpus was screened directly by the agent (coverage "
    f"{pct(SM['coverage'], 0)}, {SM['row_count']:,} rows). Screening quality was cross-checked "
    f"against an AI reference standard scored blind on a stratified sample of "
    f"{REF['sample']['sample_size']} rows across {ROUNDS['count']} rounds. Weighted "
    f"sensitivity_vs_ai_reference was {num(WM['sensitivity_vs_ai_reference'])} and "
    f"specificity_vs_ai_reference was {num(WM['specificity_vs_ai_reference'])}. After "
    f"Rogan–Gladen correction the corpus-level retain volume falls from "
    f"{CORPUS['apparent_retain_count']:,} to "
    f"{round(CORPUS['rogan_gladen_corrected_retain_count']):,} rows. The reference standard is "
    "AI-generated, not a human gold standard, so these values are cross-check results rather "
    "than clinical accuracy. No human research decision was made at any stage."
)

doc.add_page_break()
doc.add_heading("목차", 1)
for heading, subs in sections:
    doc.add_paragraph(heading, style="Heading 2")
    for sub, _ in subs:
        if sub:
            doc.add_paragraph(sub)
doc.add_paragraph("참고문헌")
doc.add_paragraph("부록 A. 질문별 결과표")
doc.add_paragraph("부록 B. 재현 경로와 파일 해시")

TABLES = {
    "1.3 연구 목적과 질문": (
        "표 1. 연구 질문의 P·I·O 구성",
        ["질문 ID", "대상(P)", "노출(I)", "결과(O)"],
        picos_rows,
    ),
    "3.2 검색 실행과 코퍼스": (
        "표 2. 질문별 검색 실행 결과",
        ["질문 ID", "대상(P)", "hit", "확보"],
        search_rows,
    ),
    "4.2 선별 결과": ("표 3. 질문별 선별 판정 분포", ["질문 ID", "retain", "deprioritize", "uncertain"], screen_rows),
    "4.4 근거 번들": ("표 4. 질문별 근거 번들 유지 행 수", ["질문 ID", "kept"], gate_rows),
    "3.4 AI 참조표준의 생성과 채점": (
        "표 5. 층화 표본 설계",
        ["층", "프레임 크기", "표본 수", "가중치"],
        stratum_rows,
    ),
    "4.3 참조 판정과의 대조": (
        "표 6. 라운드 간 일치",
        ["비교", "Cohen kappa", "라벨 일치율", "상이 축 셀"],
        round_rows,
    ),
}

for index, (heading, subs) in enumerate(sections):
    if index == 0:
        doc.add_page_break()
    doc.add_heading(heading, 1)
    for sub, paragraphs in subs:
        if sub:
            doc.add_heading(sub, 2)
        for text in paragraphs:
            doc.add_paragraph(text)
        if sub in TABLES:
            caption, headers, rows = TABLES[sub]
            doc.add_paragraph(caption)
            table(doc, headers, rows)
        if sub == "4.3 참조 판정과의 대조":
            doc.add_paragraph("표 7. AI 참조표준 대비 지표와 95% 신뢰구간")
            table(doc, ["지표", "값", "95% CI"], metric_rows)

doc.add_heading("참고문헌", 1)
REFS = [
    "Page MJ, et al. The PRISMA 2020 statement: an updated guideline for reporting systematic reviews. BMJ. 2021;372:n71.",
    "Rethlefsen ML, et al. PRISMA-S: an extension to the PRISMA Statement for Reporting Literature Searches in Systematic Reviews. Syst Rev. 2021;10:39.",
    "Miake-Lye IM, et al. What is an evidence map? A systematic review of published evidence maps and their definitions, methods, and products. Syst Rev. 2016;5:28.",
    "Rogan WJ, Gladen B. Estimating prevalence from the results of a screening test. Am J Epidemiol. 1978;107:71-76.",
    "Efron B, Tibshirani RJ. An Introduction to the Bootstrap. Chapman & Hall; 1993.",
    "Landis JR, Koch GG. The measurement of observer agreement for categorical data. Biometrics. 1977;33:159-174.",
    "National Library of Medicine. PubMed and NCBI E-utilities documentation.",
    "연구 프로토콜 v3.0. 완전 AI 자율 근거지도 트랙. research/protocol/protocol-v3.0-full-ai.md.",
]
for ref in REFS:
    doc.add_paragraph(ref, style="List Number")

doc.add_heading("부록 A. 질문별 결과표", 1)
table(
    doc,
    ["질문 ID", "hit", "확보", "retain", "deprioritize", "uncertain", "번들 kept"],
    [
        [
            qid,
            f"{RUNS[qid]['hit_count']:,}",
            f"{RUNS[qid]['retrieved_count']:,}",
            f"{SM['distribution_by_question'][qid]['retain']:,}",
            f"{SM['distribution_by_question'][qid]['deprioritize']:,}",
            f"{SM['distribution_by_question'][qid]['uncertain']:,}",
            f"{RM['by_question'][qid]:,}",
        ]
        for qid in QORDER
    ],
)

doc.add_heading("부록 B. 재현 경로와 파일 해시", 1)
table(
    doc,
    ["항목", "상대 경로", "SHA-256"],
    [[key, path.relative_to(ROOT).as_posix(), sha(path)] for key, path in P.items()],
)
doc.add_paragraph(
    "표의 수치는 문서에 고정 입력하지 않고 각 manifest에서 읽어 생성하였다. "
    "해시가 달라지면 검증 절차를 다시 수행해야 한다. 위 표의 SHA-256은 원시 바이트 기준이며, "
    "선별 manifest는 개행 차이를 흡수하기 위해 LF 정규화 후 해싱한 값을 따로 기록한다"
    f"(hash_method: {SM['hash_method']})."
)

doc.save(DOCX)

# ---------------------------------------------------------------- Markdown
md = [f"# {TITLE_KO.replace(chr(10), ' ')}", "", "## 국문초록", "", abstract, ""]
for heading, subs in sections:
    md += [f"## {heading}", ""]
    for sub, paragraphs in subs:
        if sub:
            md += [f"### {sub}", ""]
        md += sum(([text, ""] for text in paragraphs), [])
MD.write_text("\n".join(md), encoding="utf-8")

print(
    json.dumps(
        {
            "docx": DOCX.relative_to(ROOT).as_posix(),
            "markdown": MD.relative_to(ROOT).as_posix(),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        },
        ensure_ascii=False,
    )
)
